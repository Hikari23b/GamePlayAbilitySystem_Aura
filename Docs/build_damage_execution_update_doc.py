# -*- coding: utf-8 -*-
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "Docs/Aura_DamageExecution_Calculation_Update_Guide.docx"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def borders(table, color="D7DEE8", size="4"):
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = tbl_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tbl_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for i, w in enumerate(widths):
            cell = row.cells[i]
            cell.width = Inches(w)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(w * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def run_font(run, size=None, bold=None, color=None, mono=False):
    run.font.name = "Consolas" if mono else "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def style(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        st = doc.styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        st = doc.styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(11)
        st.paragraph_format.left_indent = Inches(0.375)
        st.paragraph_format.first_line_indent = Inches(-0.188)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.25


def title(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("Aura 伤害执行计算与扣血流程更新技术文档")
    run_font(r, 22, True, "0B2545")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("SetByCaller Damage / ExecCalc_Damage / Block / Armor / Critical / IncomingDamage / Floating Text")
    run_font(r, 11, False, "555555")

    rows = [
        ("项目路径", r"F:\ueprojiect\Aura"),
        ("本次主题", "火球伤害从 SetByCaller 进入 GE_Damage，经 ExecCalc_Damage 计算格挡、护甲、暴击和抗性，最后输出到 IncomingDamage 扣血。"),
        ("承接文档", "Aura_EnemyClass_DefaultAttributes_Update_Guide.docx"),
        ("关键代码", r"Source\Aura\Private\AbilitySystem\ExecCalc\ExecCalc_Damage.cpp"),
        ("核心结论", "ExecCalc 负责算最终伤害，AttributeSet 负责把 IncomingDamage 转成 Health 扣减、死亡、HitReact 和飘字。"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    borders(table)
    width(table, [1.7, 4.8])
    for row, (k, v) in zip(table.rows, rows):
        shade(row.cells[0], "E8EEF5")
        row.cells[0].paragraphs[0].add_run(k).bold = True
        row.cells[1].paragraphs[0].add_run(v)


def callout(doc, head, body):
    table = doc.add_table(rows=1, cols=1)
    borders(table, "B8C7D9")
    width(table, [6.5])
    cell = table.cell(0, 0)
    shade(cell, "F4F6F9")
    p = cell.paragraphs[0]
    r = p.add_run(head)
    run_font(r, bold=True, color="1F3A5F")
    p.add_run("  " + body)
    doc.add_paragraph()


def kv(doc, rows, headers=("位置 / 概念", "作用"), widths=(1.95, 4.55)):
    table = doc.add_table(rows=1, cols=2)
    borders(table)
    width(table, list(widths))
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        shade(cell, "E8EEF5")
        for r in cell.paragraphs[0].runs:
            r.bold = True
    for a, b in rows:
        cells = table.add_row().cells
        cells[0].text = a
        cells[1].text = b
        for cell in cells:
            margins(cell)
    doc.add_paragraph()


def tri(doc, headers, rows, widths=(1.45, 2.35, 2.7)):
    table = doc.add_table(rows=1, cols=3)
    borders(table)
    width(table, list(widths))
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        shade(cell, "E8EEF5")
        for r in cell.paragraphs[0].runs:
            r.bold = True
    for row_data in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row_data):
            cells[i].text = text
            margins(cells[i])
    doc.add_paragraph()


def flow(doc, heading, rows):
    doc.add_heading(heading, level=3)
    table = doc.add_table(rows=1, cols=3)
    borders(table)
    width(table, [0.7, 2.3, 3.5])
    for i, h in enumerate(("步", "发生位置", "意义")):
        cell = table.rows[0].cells[i]
        cell.text = h
        shade(cell, "E8EEF5")
        for r in cell.paragraphs[0].runs:
            r.bold = True
    for idx, (where, why) in enumerate(rows, 1):
        cells = table.add_row().cells
        cells[0].text = str(idx)
        cells[1].text = where
        cells[2].text = why
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for cell in cells:
            margins(cell)
    doc.add_paragraph()


def code(doc, text):
    table = doc.add_table(rows=1, cols=1)
    borders(table, "CBD5E1")
    width(table, [6.5])
    cell = table.cell(0, 0)
    shade(cell, "F8FAFC")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    run_font(r, 8.5, False, None, True)
    doc.add_paragraph()


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def build():
    doc = Document()
    style(doc)
    title(doc)

    doc.add_heading("1. 本次更新解决了什么", level=1)
    callout(
        doc,
        "一句话",
        "伤害不再只是固定扣血，而是先把基础伤害放进 GameplayEffectSpec，再由 ExecCalc_Damage 根据攻防双方属性、曲线表和随机判定算出最终伤害。",
    )
    tri(
        doc,
        ("阶段", "之前", "现在"),
        [
            ("伤害来源", "火球或 GE 里直接给一个固定值。", "AuraProjectileSpell 用 SetByCaller 把 Damage 写入 Spec。"),
            ("伤害计算", "AttributeSet 只接收 IncomingDamage 扣血。", "ExecCalc_Damage 读取 Source/Target 属性，计算格挡、护甲和暴击。"),
            ("防御参与", "Armor / BlockChance 等属性还没有真正影响伤害。", "Target BlockChance 可让伤害减半；Armor 按曲线降低伤害。"),
            ("攻击参与", "Source 的穿甲和暴击属性还没有进入公式。", "Source ArmorPenetration、CriticalHitChance、CriticalHitDamage 进入最终伤害。"),
            ("结果承接", "Health 直接被 GE 修改。", "ExecCalc 输出 IncomingDamage，AttributeSet 再清零、扣 Health、死亡/HitReact/飘字。"),
        ],
    )

    doc.add_heading("2. 当前文件职责地图", level=1)
    kv(
        doc,
        [
            ("AuraProjectileSpell.cpp", "服务端生成 Projectile，并创建 DamageEffectSpecHandle。通过 AssignTagSetByCallerMagnitude 写入 Damage。"),
            ("AuraProjectile.cpp", "Projectile 命中时，在服务端把 DamageEffectSpecHandle 应用到目标 ASC。"),
            ("ExecCalc_Damage.cpp", "GE_Damage 的 ExecutionCalculation。读取 SetByCaller Damage，捕获双方属性，输出 IncomingDamage。"),
            ("AuraAttributeSet.cpp", "PostGameplayEffectExecute 接收 IncomingDamage，清零中转属性，扣 Health，触发死亡/HitReact/飘字。"),
            ("CharacterClassInfo.h", "持有 DamageCalculationCoefficients CurveTable，为护甲、穿甲、暴击抗性等公式提供等级系数。"),
            ("AuraGameplayTags.cpp/.h", "注册 Damage Tag，作为 SetByCaller 的 Data Tag。"),
        ],
    )

    doc.add_heading("3. 总流程：从火球到扣血", level=1)
    flow(
        doc,
        "伤害执行链路",
        [
            ("AuraProjectileSpell::SpawnProjectile", "只在服务端生成火球，创建 DamageEffectSpecHandle。"),
            ("AssignTagSetByCallerMagnitude", "把基础伤害写入 Spec：Damage Tag -> ScaledDamage。"),
            ("AAuraProjectile::OnSphereOverlap", "命中目标时，服务端取目标 ASC，ApplyGameplayEffectSpecToSelf。"),
            ("GE_Damage", "执行 ExecCalc_Damage，而不是普通 Modifier 直接改 Health。"),
            ("ExecCalc_Damage", "从 Spec 读 Damage，从 Source/Target 捕获战斗属性，套公式后输出 IncomingDamage。"),
            ("AuraAttributeSet::PostGameplayEffectExecute", "检测 IncomingDamage，保存局部伤害后清零，再扣 Health。"),
            ("死亡 / HitReact / 飘字", "Health <= 0 调 Die；未死亡则尝试激活 HitReact；最后调用 ShowFloatingText。"),
        ],
    )

    doc.add_heading("4. SetByCaller Damage", level=1)
    code(
        doc,
        """const FGameplayEffectSpecHandle SpecHandle =
    SourceASC->MakeOutgoingSpec(DamageEffectClass, GetAbilityLevel(), SourceASC->MakeEffectContext());

const float ScaledDamage = Damage.GetValueAtLevel(10);
UAbilitySystemBlueprintLibrary::AssignTagSetByCallerMagnitude(
    SpecHandle,
    GameplayTags.Damage,
    ScaledDamage
);""",
    )
    callout(
        doc,
        "怎么理解",
        "Damage 不是 AttributeSet 上的长期属性，而是这一次攻击写进 Spec 的临时参数。ExecCalc 用 Spec.GetSetByCallerMagnitude 取它，所以同一个 GE_Damage 可以承载不同技能、不同等级的伤害。",
    )
    bullets(
        doc,
        [
            "GE_Damage 里的 Modifier 或 Execution 必须使用同一个 Data Tag，也就是 FAuraGameplayTags::Get().Damage。",
            "当前代码里 Damage.GetValueAtLevel(10) 仍是硬编码等级 10；正式做法通常换成 GetAbilityLevel()。",
            "如果 GameplayTag 选择器受 SetByCaller 过滤影响，建议 Damage Tag 命名到 SetByCaller.Damage 一类的路径。当前代码使用的是 Damage。",
        ],
    )

    doc.add_heading("5. ExecCalc 捕获属性设计", level=1)
    tri(
        doc,
        ("属性", "捕获方", "用途"),
        [
            ("Armor", "Target", "目标护甲，进入 EffectiveArmor，最终降低伤害。"),
            ("BlockChance", "Target", "目标格挡概率，随机成功后伤害减半。"),
            ("ArmorPenetration", "Source", "攻击者穿甲，降低目标有效护甲。"),
            ("CriticalHitChance", "Source", "攻击者暴击率，减去目标抗性后决定是否暴击。"),
            ("CriticalHitDamage", "Source", "攻击者暴击附加伤害，暴击时加到双倍伤害之后。"),
            ("CriticalHitResistance", "Target", "目标暴击抗性，按曲线系数削减攻击者暴击率。"),
        ],
    )
    code(
        doc,
        """DEFINE_ATTRIBUTE_CAPTUREDEF(UAuraAttributeSet, Armor, Target, false);
DEFINE_ATTRIBUTE_CAPTUREDEF(UAuraAttributeSet, BlockChance, Target, false);
DEFINE_ATTRIBUTE_CAPTUREDEF(UAuraAttributeSet, ArmorPenetration, Source, false);
DEFINE_ATTRIBUTE_CAPTUREDEF(UAuraAttributeSet, CriticalHitChance, Source, false);
DEFINE_ATTRIBUTE_CAPTUREDEF(UAuraAttributeSet, CriticalHitDamage, Source, false);
DEFINE_ATTRIBUTE_CAPTUREDEF(UAuraAttributeSet, CriticalHitResistance, Target, false);""",
    )
    callout(
        doc,
        "Source 和 Target",
        "Source 是攻击者，Target 是被打者。玩家火球打敌人时，CriticalHitChance 读玩家；敌人攻击玩家时，CriticalHitChance 读敌人。这和 Secondary GE 是否共享不是同一个问题。",
    )

    doc.add_heading("6. 当前伤害公式拆解", level=1)
    flow(
        doc,
        "ExecCalc_Damage 内部计算顺序",
        [
            ("读取 SetByCaller Damage", "Spec.GetSetByCallerMagnitude(Damage)，得到技能塞进来的基础伤害。"),
            ("格挡判定", "读取 TargetBlockChance，随机 1 到 100，小于 BlockChance 时 Damage / 2。"),
            ("读取护甲和穿甲", "TargetArmor 来自目标，SourceArmorPenetration 来自攻击者。"),
            ("读取暴击三件套", "SourceCriticalHitChance、SourceCriticalHitDamage、TargetCriticalHitResistance。"),
            ("暴击抗性曲线", "从 DamageCalculationCoefficients 查 CriticalHitResistance 行，用目标等级 Eval 系数。"),
            ("有效暴击率", "EffectiveCriticalHitChance = SourceCritChance - TargetCritResistance * Coefficient。"),
            ("暴击判定", "随机成功时 Damage = Damage * 2 + SourceCriticalHitDamage。"),
            ("穿甲曲线", "查 ArmorPenetration 行，用攻击者等级 Eval 系数。"),
            ("有效护甲", "EffectiveArmor = TargetArmor * (100 - SourceArmorPenetration * Coefficient) / 100。"),
            ("护甲减伤曲线", "查 EffectiveArmor 行，用目标等级 Eval 系数。"),
            ("最终输出", "Damage *= (100 - EffectiveArmor * Coefficient) / 100，然后 IncomingDamage += Damage。"),
        ],
    )
    code(
        doc,
        """const FGameplayModifierEvaluatedData EvaluatedData(
    UAuraAttributeSet::GetIncomingDamageAttribute(),
    EGameplayModOp::Additive,
    Damage
);
OutExecutionOutput.AddOutputModifier(EvaluatedData);""",
    )

    doc.add_heading("7. CurveTable 系数表", level=1)
    tri(
        doc,
        ("Row Name", "当前用途", "注意"),
        [
            ("ArmorPenetration", "按 Source 等级计算穿甲系数。", "C++ FindCurve 名字必须完全一致。"),
            ("EffectiveArmor", "按 Target 等级计算护甲减伤系数。", "行缺失时 FindCurve 返回 nullptr，直接 Eval 会崩。"),
            ("CriticalHitResistance", "按 Target 等级计算暴击抗性系数。", "如果表里没有这一行，当前 ExecCalc 会空指针崩溃。"),
        ],
    )
    callout(
        doc,
        "为什么放 CharacterClassInfo",
        "DamageCalculationCoefficients 是整套战斗公式的公共配置，不属于某个技能或某个敌人实例。GameMode 持有 CharacterClassInfo 后，服务端 ExecCalc 可以通过 Library 拿到当前关卡使用的公式系数表。",
    )

    doc.add_heading("8. IncomingDamage 扣血与后续效果", level=1)
    code(
        doc,
        """if (Data.EvaluatedData.Attribute == GetIncomingDamageAttribute())
{
    const float LocalIncomingDamage = GetIncomingDamage();
    SetIncomingDamage(0.f);

    const float NewHealth = GetHealth() - LocalIncomingDamage;
    SetHealth(FMath::Clamp(NewHealth, 0.f, GetMaxHealth()));
}""",
    )
    tri(
        doc,
        ("步骤", "代码行为", "目的"),
        [
            ("读取", "LocalIncomingDamage = GetIncomingDamage()", "把 ExecCalc 输出的最终伤害取出来。"),
            ("清零", "SetIncomingDamage(0.f)", "IncomingDamage 是 Meta Attribute，用完就清，防止重复结算。"),
            ("扣血", "Health = Clamp(Health - Damage, 0, MaxHealth)", "真正改变生命值。"),
            ("死亡", "NewHealth <= 0 时 Cast ICombatInterface 并调用 Die。", "死亡逻辑交给角色实现。"),
            ("受击", "未死亡时按 Effects_HitReact 激活受击能力。", "让敌人/角色播放受击反应。"),
            ("飘字", "ShowFloatingText(Props, LocalIncomingDamage)", "通过 Source 玩家控制器显示伤害数字。"),
        ],
    )

    doc.add_heading("9. 网络职责与 GameMode 注意点", level=1)
    callout(
        doc,
        "关键坑",
        "GameMode 只存在于服务器。ExecCalc_Damage 通过 GetCharacterClassInfo(SourceAvatar) 取 GameMode 上的 CharacterClassInfo，因此这套伤害计算必须由服务端权威执行。客户端如果也跑进 ExecCalc，会拿不到 GameMode 而空指针。",
    )
    bullets(
        doc,
        [
            "Projectile 的 ApplyGameplayEffectSpecToSelf 当前放在 HasAuthority() 内，这是正确方向。",
            "如果客户端仍然崩，优先检查是否有客户端路径也应用了 GE_Damage。",
            "CharacterClassInfo、DamageCalculationCoefficients、每条 Curve 行都建议在调试期用 check 或日志验证。",
            "SourceCombatInterface 和 TargetCombatInterface 当前直接使用 GetPlayerLevel，实际项目里也应该判空或保证双方都实现 CombatInterface。",
        ],
    )

    doc.add_heading("10. 当前实现里值得记账的点", level=1)
    bullets(
        doc,
        [
            "格挡和暴击判定目前使用 <，如果想让 50 表示完整 50% 概率，通常写 <= 更直观。",
            "EffectiveArmor 当前写成 TargetArmor *= ...，会同时修改 TargetArmor 自身；虽然结果可用，但单独写 const float EffectiveArmor = TargetArmor * ... 更清楚。",
            "AuraProjectileSpell 当前 Damage.GetValueAtLevel(10) 硬编码等级 10，后续建议改为 GetAbilityLevel()。",
            "LoopingSoundComponent->Stop() 当前未判空，如果 LoopingSound 未配置，Projectile 命中或销毁时可能崩。",
            "玩家和敌人的 Secondary 属性初始化路径不同：敌人走 CharacterClassInfo，玩家仍走 CharacterBase 的 DefaultSecondaryAttributes。调试 Source 属性时先确认 Source 是谁。",
        ],
    )

    doc.add_heading("11. 心智模型总结", level=1)
    callout(
        doc,
        "Spec 是伤害快递单",
        "火球生成时把 Damage 写在 Spec 里，Projectile 只是把这张快递单带到目标身上。ExecCalc 打开快递单，读出基础伤害。",
    )
    callout(
        doc,
        "ExecCalc 是结算台",
        "结算台会看攻击者的穿甲和暴击，也看目标的护甲、格挡和抗性，最后算出真正要扣多少。",
    )
    callout(
        doc,
        "IncomingDamage 是中转站",
        "ExecCalc 不直接改 Health，而是把最终伤害写入 IncomingDamage。AttributeSet 看到中转站有货，取出来扣血，然后立刻清空。"
    )

    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
