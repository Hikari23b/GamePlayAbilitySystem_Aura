# -*- coding: utf-8 -*-
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "Docs/Aura_EnemyHealthBar_Projectile_Update_Guide.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="B8C7D9", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def apply_run_font(run, size=None, bold=None, color=None, mono=False):
    run.font.name = "Consolas" if mono else "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_title(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("Aura 敌人血条、鼠标目标数据与火球伤害更新技术文档")
    apply_run_font(run, size=22, bold=True, color="0B2545")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("对照上一版 InputTag / AutoRun / Projectile 文档整理：Enemy HealthBar / TargetDataHandle / DamageEffectSpecHandle")
    apply_run_font(run, size=11, color="555555")

    table = doc.add_table(rows=5, cols=2)
    set_table_borders(table, "D7DEE8")
    set_table_width(table, [1.7, 4.8])
    rows = [
        ("项目路径", r"F:\ueprojiect\Aura"),
        ("文档定位", "记录你在上一版之后完成的敌人头顶血条、WidgetController 注入、火球伤害 Spec 传递、鼠标目标数据 AbilityTask 等改动。"),
        ("核心变化", "敌人现在既是战斗单位，也是自己头顶血条 Widget 的 Controller；火球从技能生成后携带 DamageEffectSpecHandle，并在命中时由服务端应用伤害。"),
        ("关联旧文档", "Aura_InputTag_AutoRun_Projectile_Technical_Guide.docx；Aura_AttributeMenu_Technical_Guide.docx。"),
        ("当前状态", "代码侧已有 AAuraEnemy HealthBar 组件与属性广播；蓝图侧新增 WBP_EnemyHealthBar、WBP_ProgressBar，并在敌人蓝图里配置显示。"),
    ]
    for row, (label, value) in zip(table.rows, rows):
        set_cell_shading(row.cells[0], "E8EEF5")
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[1].paragraphs[0].add_run(value)


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_borders(table, "B8C7D9")
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    r = p.add_run(title)
    apply_run_font(r, bold=True, color="1F3A5F")
    p.add_run("  " + body)
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_code(doc, code):
    table = doc.add_table(rows=1, cols=1)
    set_table_borders(table, "CBD5E1")
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F8FAFC")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(code)
    apply_run_font(run, size=8.5, mono=True)
    doc.add_paragraph()


def add_kv_table(doc, rows, widths=(1.9, 4.6), headers=("位置 / 名称", "作用")):
    table = doc.add_table(rows=1, cols=2)
    set_table_borders(table, "D7DEE8")
    set_table_width(table, list(widths))
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, "E8EEF5")
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
        for cell in cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()


def add_three_col_table(doc, headers, rows, widths=(1.45, 2.35, 2.7)):
    table = doc.add_table(rows=1, cols=3)
    set_table_borders(table, "D7DEE8")
    set_table_width(table, list(widths))
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, "E8EEF5")
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row_data in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row_data):
            cells[i].text = text
            set_cell_margins(cells[i])
    doc.add_paragraph()


def add_flow_table(doc, title, steps):
    doc.add_heading(title, level=3)
    table = doc.add_table(rows=1, cols=3)
    set_table_borders(table, "D7DEE8")
    set_table_width(table, [0.7, 2.35, 3.45])
    for i, header in enumerate(("步", "发生位置", "意义")):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, "E8EEF5")
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for idx, (where, meaning) in enumerate(steps, 1):
        cells = table.add_row().cells
        cells[0].text = str(idx)
        cells[1].text = where
        cells[2].text = meaning
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for cell in cells:
            set_cell_margins(cell)
    doc.add_paragraph()


def build_doc():
    doc = Document()
    style_doc(doc)
    add_title(doc)

    doc.add_heading("1. 本次相对上一版的新增点", level=1)
    add_callout(
        doc,
        "一句话总览",
        "上一版重点是输入 Tag、自动寻路和火球生成；这一版补上了敌人头顶血条，以及火球命中后真正把 GameplayEffectSpec 应用到目标 ASC 的链路。",
    )
    add_three_col_table(
        doc,
        ("模块", "新增 / 改动", "现在要理解的关键点"),
        [
            ("AAuraEnemy", "新增 UWidgetComponent HealthBar；新增 OnHealthChanged、OnMaxHealthChanged 蓝图可绑定委托。", "敌人自己充当头顶血条 WidgetController，不再走 HUD 那套全局 UI Controller。"),
            ("WBP_EnemyHealthBar", "作为 WidgetComponent 的 WidgetClass，内部再组合 WBP_ProgressBar。", "WidgetControllerSet 里 Cast 到敌人，绑定敌人广播出来的 Health / MaxHealth。"),
            ("UTargetDataUnderMouse", "ValidData 参数改成 FGameplayAbilityTargetDataHandle。", "更贴合 GAS 的目标数据复制模型，客户端本地取鼠标命中，服务端通过 replicated target data 接收。"),
            ("UAuraProjectileSpell", "SpawnProjectile 通过 Deferred Spawn 创建 Projectile，并给 Projectile 写入 DamageEffectSpecHandle。", "火球 Actor 不自己凭空构造伤害，而是带着技能创建好的 Spec 去命中目标。"),
            ("AAuraProjectile", "命中时播放 Impact 表现、停止 LoopingSound、服务端 ApplyGameplayEffectSpecToSelf。", "表现可在客户端兜底播放，伤害只在服务端发生。"),
            ("Aura.Build.cs / Aura.h", "加入 NavigationSystem、Niagara 等依赖，定义 ECC_Projectile。", "寻路、特效、Projectile 碰撞通道都需要模块或自定义 Trace Channel 支撑。"),
        ],
    )

    doc.add_heading("2. 当前文件职责地图", level=1)
    add_kv_table(
        doc,
        [
            ("Source/Aura/Public/Character/AuraEnemy.h", "声明敌人接口实现、等级、头顶血条组件、Health / MaxHealth 两个动态多播委托。"),
            ("Source/Aura/Private/Character/AuraEnemy.cpp", "构造 HealthBar 组件；BeginPlay 初始化 ASC、给 Widget 注入 Controller、绑定属性变化并广播初始值。"),
            ("Source/Aura/Public/UI/Widget/AuraUserWidget.h", "提供通用 WidgetController 字段和 SetWidgetController 入口，蓝图通过 WidgetControllerSet 事件继续绑定。"),
            ("Source/Aura/Public/AbilitySystem/AbilityTasks/TargetDataUnderMouse.h", "声明自定义 AbilityTask 节点和 ValidData 输出引脚。"),
            ("Source/Aura/Private/AbilitySystem/AbilityTasks/TargetDataUnderMouse.cpp", "客户端取鼠标命中，服务端等待 TargetData 复制回调，并统一广播 ValidData。"),
            ("Source/Aura/Private/AbilitySystem/Abilities/AuraProjectileSpell.cpp", "在服务端从战斗 Socket 生成火球，并把 DamageEffectSpecHandle 塞给 Projectile。"),
            ("Source/Aura/Private/Actor/AuraProjectile.cpp", "设置 Projectile 碰撞和移动；BeginPlay 绑定 Overlap；命中时应用伤害、播放特效和销毁。"),
            ("Content/BluePrints/UI/ProgressBar/WBP_EnemyHealthBar.uasset", "敌人头顶血条外层 Widget，接收敌人 Controller 并更新血条百分比。"),
            ("Content/BluePrints/UI/ProgressBar/WBP_ProgressBar.uasset", "通用进度条组件，可做真实血条和 Ghost Bar 插值。"),
        ],
    )

    doc.add_heading("3. 敌人头顶血条的完整流程", level=1)
    add_flow_table(
        doc,
        "从敌人生成到血条显示",
        [
            ("AAuraEnemy 构造函数", "CreateDefaultSubobject<UWidgetComponent>(\"HealthBar\") 创建组件，并挂到 RootComponent。"),
            ("BP_EnemyBase / 子蓝图", "在组件 Details 里设置 WidgetClass = WBP_EnemyHealthBar，设置 Space、Draw Size、Relative Location。"),
            ("AAuraEnemy::BeginPlay", "先 InitAbilityActorInfo 和 InitializeDefaultAttributes，保证 AttributeSet 里有 Health / MaxHealth 初始值。"),
            ("HealthBar->GetUserWidgetObject()", "拿到 WidgetComponent 内部实际创建出来的 UUserWidget 实例。"),
            ("AuraUserWidget->SetWidgetController(this)", "把当前 AAuraEnemy 指针塞进 WidgetController，然后触发蓝图 Event WidgetControllerSet。"),
            ("WBP_EnemyHealthBar 蓝图", "在 WidgetControllerSet 中 Cast 到 AAuraEnemy，绑定 OnHealthChanged 和 OnMaxHealthChanged。"),
            ("ASC 属性变化委托", "C++ 绑定 GetGameplayAttributeValueChangeDelegate，属性变化时广播新值给 Widget。"),
            ("初始广播", "BeginPlay 末尾主动 Broadcast 当前 Health / MaxHealth，避免等第一次受伤时 UI 才有数据。"),
        ],
    )
    add_code(
        doc,
        """if (UAuraUserWidget* AuraUserWidget = Cast<UAuraUserWidget>(HealthBar->GetUserWidgetObject()))
{
    AuraUserWidget->SetWidgetController(this);
}

AbilitySystemComponent->GetGameplayAttributeValueChangeDelegate(AuraAS->GetHealthAttribute()).AddLambda(
    [this](const FOnAttributeChangeData& Data)
    {
        OnHealthChanged.Broadcast(Data.NewValue);
    }
);

OnHealthChanged.Broadcast(AuraAS->GetHealth());
OnMaxHealthChanged.Broadcast(AuraAS->GetMaxHealth());""",
    )
    add_callout(
        doc,
        "这里为什么让敌人自己当 Controller",
        "敌人头顶血条是每个敌人自己的局部 UI，不需要通过 PlayerController 或 HUD 全局查找。谁拥有这条血，谁广播这条血，Widget 只要绑定这个敌人即可。",
    )

    doc.add_heading("4. 蓝图侧 WBP_EnemyHealthBar / WBP_ProgressBar 检查表", level=1)
    add_three_col_table(
        doc,
        ("检查项", "正确姿势", "常见问题"),
        [
            ("WidgetComponent", "WidgetClass 指向 WBP_EnemyHealthBar；Space 推荐 Screen；Draw Size 如 120x20；Z 轴抬到头顶。", "WidgetClass 没设或 Draw Size 太小，会导致游戏里看不到。"),
            ("父类", "WBP_EnemyHealthBar 最好继承 UAuraUserWidget 或其蓝图父类。", "如果不是 AuraUserWidget，C++ Cast 会失败，SetWidgetController 不执行。"),
            ("WidgetControllerSet", "在该事件里 Cast WidgetController 到 AAuraEnemy，再 Bind Event 到 OnHealthChanged / OnMaxHealthChanged。", "只写 Construct 不够，因为 Controller 是 C++ 后注入的。"),
            ("Percent 计算", "保存 CurrentHealth 和 CurrentMaxHealth，ProgressBar Percent = Health / MaxHealth。", "MaxHealth 为 0 或尚未收到初始广播，会导致 Percent 异常。"),
            ("ProgressBarVisible", "它只管显示隐藏，不建议拿来控制 Ghost Bar 插值。", "Construct 里 Set Bar Visibility 输入不接线，会默认传 false，把运行时可见性关掉。"),
            ("Ghost Bar", "单独用 bInterpGhostBar 或目标百分比控制插值。", "把可见性变量当插值开关，会出现 true 不插值、false 才插值的混乱。"),
        ],
    )

    doc.add_heading("5. 输入、鼠标目标数据与技能激活链路", level=1)
    add_flow_table(
        doc,
        "从按键到 Ability",
        [
            ("UAuraInputComponent::BindAbilityActions", "遍历 InputConfig，把每个 InputAction 的 Started / Completed / Triggered 绑定到 Pressed / Released / Held。"),
            ("AAuraPlayerController::AbilityInputTagHeld", "非 LMB 直接转发给 ASC；LMB 如果正在选敌人或按着 Shift，也转发给 ASC。"),
            ("UAuraAbilitySystemComponent::AbilityInputTagHeld", "遍历 ActivatableAbilities，用 DynamicAbilityTags 匹配 InputTag，匹配到就 TryActivateAbility。"),
            ("UAuraGameplayAbility::StartupInputTag", "技能蓝图在默认值里配置自己的输入 Tag，GiveAbility 时写进 AbilitySpec.DynamicAbilityTags。"),
            ("GA_FireBolt 蓝图", "激活后可以调用 TargetDataUnderMouse，等 ValidData 后再调用 SpawnProjectile。"),
        ],
    )
    add_code(
        doc,
        """for (FGameplayAbilitySpec& AbilitySpec : GetActivatableAbilities())
{
    if (AbilitySpec.DynamicAbilityTags.HasTagExact(InputTag))
    {
        AbilitySpecInputPressed(AbilitySpec);
        if (!AbilitySpec.IsActive())
        {
            TryActivateAbility(AbilitySpec.Handle);
        }
    }
}""",
    )

    doc.add_heading("6. TargetDataUnderMouse 的 GAS 目标数据流程", level=1)
    add_flow_table(
        doc,
        "客户端预测与服务端接收",
        [
            ("CreateTargetDataUnderMouse", "蓝图节点创建 AbilityTask，因此 ValidData 委托会变成该节点的输出执行引脚。"),
            ("Activate", "判断当前 Ability 是否本地控制。本地控制端负责取鼠标命中；非本地端等待复制数据。"),
            ("SendMouseCursorData", "GetHitResultUnderCursor 得到 CursorHit，并包装成 FGameplayAbilityTargetData_SingleTargetHit。"),
            ("ServerSetReplicatedTargetData", "把 DataHandle 发给服务端，同时带上 AbilitySpecHandle 和 PredictionKey。"),
            ("OnTargetDataReplicatedCallBack", "服务端收到后 ConsumeClientReplicatedTargetData，再 Broadcast ValidData。"),
            ("ValidData", "蓝图从 DataHandle 里 Break 出 HitResult / Location，用这个位置决定火球朝向或运动扭曲目标。"),
        ],
    )
    add_callout(
        doc,
        "为什么不用 FVector 直接传",
        "FGameplayAbilityTargetDataHandle 是 GAS 原生目标数据容器，能和预测、复制、AbilityTask 输出引脚配合。以后要从点目标升级到 Actor 目标、HitResult、范围目标，也更顺。",
    )

    doc.add_heading("7. 火球生成、Spec 传递与命中伤害", level=1)
    add_flow_table(
        doc,
        "从 SpawnProjectile 到 ApplyGameplayEffect",
        [
            ("UAuraProjectileSpell::SpawnProjectile", "只在服务端执行，防止客户端自己生成有伤害权限的火球。"),
            ("CombatInterface->GetCombatSocketLocation", "从角色武器 Socket 取发射位置。玩家和敌人都可以通过 CharacterBase 的实现复用。"),
            ("SpawnActorDeferred", "先创建未完成初始化的 Projectile，给它塞必要数据，再 FinishSpawning。"),
            ("MakeOutgoingSpec", "用 DamageEffectClass 和技能等级创建 FGameplayEffectSpecHandle。"),
            ("Projectile->DamageEffectSpecHandle", "把伤害 Spec 交给火球保存，火球命中时使用。"),
            ("AAuraProjectile::OnSphereOverlap", "命中后找 OtherActor 的 ASC，服务端 ApplyGameplayEffectSpecToSelf，然后 Destroy。"),
            ("AAuraProjectile::Destroyed", "客户端如果没走到 Overlap 表现，用 Destroyed 做一次 Impact 音效和 Niagara 兜底。"),
        ],
    )
    add_code(
        doc,
        """AAuraProjectile* Projectile = GetWorld()->SpawnActorDeferred<AAuraProjectile>(
    ProjectileClass,
    SpawnTransform,
    GetOwningActorFromActorInfo(),
    Cast<APawn>(GetOwningActorFromActorInfo()),
    ESpawnActorCollisionHandlingMethod::AlwaysSpawn);

const FGameplayEffectSpecHandle SpecHandle =
    SourceASC->MakeOutgoingSpec(DamageEffectClass, GetAbilityLevel(), SourceASC->MakeEffectContext());

Projectile->DamageEffectSpecHandle = SpecHandle;
Projectile->FinishSpawning(SpawnTransform);""",
    )
    add_callout(
        doc,
        "Deferred Spawn 的意义",
        "它像先把火球演员叫到后台化妆：Actor 已经有了，但还没正式登场。你可以先把伤害 Spec、Owner、Instigator、朝向等关键数据塞好，再 FinishSpawning 让它进入世界。",
    )

    doc.add_heading("8. 网络职责：谁负责表现，谁负责真实结果", level=1)
    add_three_col_table(
        doc,
        ("事情", "客户端", "服务端"),
        [
            ("输入", "本地立刻响应按键、取鼠标命中、做预测。", "接收 Input / TargetData，判断技能能否真正生效。"),
            ("寻路", "PlayerController 本地控制移动和 AutoRun 体验。", "多人项目里还需要考虑权威移动与同步。"),
            ("火球生成", "不应该生成有伤害权限的正式 Projectile。", "SpawnProjectile 通过 HasAuthority 限制只在服务端生成。"),
            ("命中特效", "可以播放 Impact 音效 / Niagara，Destroyed 做兜底。", "也播放表现，并负责销毁 replicated Projectile。"),
            ("伤害", "不直接 Apply GameplayEffect。", "通过 TargetASC->ApplyGameplayEffectSpecToSelf 应用伤害。"),
            ("敌人血条", "显示收到的 Health / MaxHealth 广播。", "AttributeSet 的真实变化来自 ASC 和 GE。"),
        ],
    )

    doc.add_heading("9. 当前代码风险点与排查建议", level=1)
    add_bullets(
        doc,
        [
            "AAuraEnemy::BeginPlay 里使用 HealthBar->GetUserWidgetObject 前，最好确认 HealthBar 不为空，并确认 WidgetClass 已在蓝图设置；否则 Cast 失败后血条不会绑定。",
            "Aura.Build.cs 当前 PublicDependencyModuleNames 没显式写 UMG。如果以后 WidgetComponent 相关链接或编译异常，建议把 UMG 加进依赖。",
            "AAuraProjectile::Destroyed 和 OnSphereOverlap 中调用 LoopingSoundComponent->Stop() 前建议判空，避免 LoopingSound 没配置时崩溃。",
            "OnSphereOverlap 目前没有过滤 OtherActor 是否是自己、Owner、友方或已经命中过的目标。后续做战斗规则时建议补充过滤和 bHit 防重复。",
            "DamageEffectSpecHandle.Data.Get() 使用前最好确认 SpecHandle 有效；如果 DamageEffectClass 没设置，命中时可能崩。",
            "敌人头顶血条如果游戏中不可见，优先查 WidgetClass、Space、DrawSize、RelativeLocation、Widget 父类、ProgressBar Percent 和 Construct 里是否把可见性设成 false。",
            "ProgressBarVisible 只建议管 Visibility；虚血条插值建议用 bInterpGhostBar 或 TargetPercent 单独控制。",
        ],
    )

    doc.add_heading("10. 推荐的下一步演进", level=1)
    add_numbered(
        doc,
        [
            "给 Projectile 命中加入目标过滤：忽略 Owner、忽略无 ASC 对象、按阵营判断敌我。",
            "把 DamageEffectSpec 的 Context 补充 SourceObject / HitResult，方便后续在 AttributeSet 或 ExecutionCalculation 里取命中信息。",
            "把敌人血条显示策略做成蓝图变量：永远显示、受伤后显示几秒、鼠标悬停显示。",
            "把 WBP_ProgressBar 的真实条、Ghost 条、插值速度整理成一个稳定父类，子类只改颜色和尺寸。",
            "把 AbilityTask 输出的 TargetData 与 Motion Warping 的 UpdateFacingTarget 串起来，让玩家和敌人的施法朝向都走 CombatInterface。",
        ],
    )

    doc.add_heading("11. 心智模型总结", level=1)
    add_callout(
        doc,
        "敌人血条",
        "敌人是数据源，HealthBar 是显示器，OnHealthChanged 是电线。BeginPlay 时先把电线插好，再广播一次初始电流，UI 才能一出现就亮。",
    )
    add_callout(
        doc,
        "火球伤害",
        "技能负责决定伤害内容，火球负责把这份伤害送到目标身上。Projectile 不应该临场瞎编伤害，它只是携带 Spec 的快递员。",
    )
    add_callout(
        doc,
        "TargetData",
        "鼠标点到哪里是客户端最清楚，但最终能不能打中和造成什么结果要让服务端知道。TargetDataHandle 就是把这个命中信息用 GAS 听得懂的格式打包寄过去。",
    )

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
