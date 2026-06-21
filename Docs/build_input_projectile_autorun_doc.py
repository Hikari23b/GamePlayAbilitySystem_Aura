from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "Docs/Aura_InputTag_AutoRun_Projectile_Technical_Guide.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="B8C7D9", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_title(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("Aura 输入 Tag、自动寻路与投射物技能技术文档")
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Enhanced Input / InputTag / AutoRun / GameplayAbility / Projectile")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string("555555")

    table = doc.add_table(rows=5, cols=2)
    set_table_borders(table, "D7DEE8")
    set_table_width(table, [1.7, 4.8])
    rows = [
        ("项目路径", "F:\\ueprojiect\\Aura"),
        ("文档定位", "继属性菜单文档之后，记录新完成的鼠标寻路、输入 Tag 激活技能、投射物生成链路"),
        ("核心主题", "左键两用：点敌人走技能输入，点地面走短按寻路或长按跟随鼠标"),
        ("相关旧文档", "Aura_AttributeMenu_Technical_Guide.docx、Aura_GameplayEffect_Technical_Guide_v2_Naming.docx"),
        ("当前代码状态", "InputTag 已注册，InputConfig 批量绑定，ASC 通过 DynamicAbilityTags 找技能，ProjectileSpell 在服务器生成 Projectile"),
    ]
    for row, (label, value) in zip(table.rows, rows):
        set_cell_shading(row.cells[0], "E8EEF5")
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[1].paragraphs[0].add_run(value)


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_borders(table, "B8C7D9")
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string("1F3A5F")
    p.add_run("  " + body)
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.add_run(item)


def add_code(doc, code):
    table = doc.add_table(rows=1, cols=1)
    set_table_borders(table, "CBD5E1")
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F8FAFC")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(code)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(8.5)


def add_kv_table(doc, rows, widths=(1.9, 4.6), headers=("位置 / 名称", "作用")):
    table = doc.add_table(rows=1, cols=2)
    set_table_borders(table, "D7DEE8")
    set_table_width(table, list(widths))
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, "E8EEF5")
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
        for cell in cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()


def add_three_col_table(doc, headers, rows, widths=(1.5, 2.35, 2.65)):
    table = doc.add_table(rows=1, cols=3)
    set_table_borders(table, "D7DEE8")
    set_table_width(table, list(widths))
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, "E8EEF5")
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row_data in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row_data):
            cells[i].text = text
            set_cell_margins(cells[i])
    doc.add_paragraph()


def build_doc():
    doc = Document()
    style_doc(doc)
    add_title(doc)

    doc.add_heading("1. 总体心智模型", level=1)
    add_callout(
        doc,
        "一句话版本",
        "这套新逻辑把玩家输入先翻译成 InputTag；左键 InputTag 会先判断鼠标下是不是敌人，决定走技能还是走移动；技能触发后由 ASC 找到匹配 Tag 的 AbilitySpec，并在服务器上的 ProjectileSpell 中生成投射物。",
    )
    add_code(
        doc,
        "鼠标 / 键盘输入\n"
        "    -> Enhanced Input 的 UInputAction\n"
        "    -> UAuraInputConfig 映射成 FGameplayTag\n"
        "    -> AAuraPlayerController 判断 LMB 是攻击还是移动\n"
        "    -> UAuraAbilitySystemComponent 用 InputTag 查 AbilitySpec\n"
        "    -> TryActivateAbility\n"
        "    -> UAuraProjectileSpell::ActivateAbility\n"
        "    -> SpawnActorDeferred<AAuraProjectile>\n"
        "    -> AAuraProjectile 通过 Sphere + ProjectileMovement 飞行和碰撞",
    )
    doc.add_paragraph(
        "这和属性菜单文档里的模式是同一类思想：属性菜单是 Tag 把“属性”和“UI 行”配对；这里是 Tag 把“输入动作”和“技能”配对。"
    )

    doc.add_heading("2. 文件职责总览", level=1)
    add_kv_table(
        doc,
        [
            ("AuraGameplayTags.h/.cpp", "注册 InputTag.LMB、InputTag.RMB、InputTag.1 到 InputTag.4，提供 FAuraGameplayTags::Get() 统一访问。"),
            ("AuraInputConfig.h/.cpp", "DataAsset：保存 UInputAction 与 InputTag 的对应关系，例如 IA_LMB -> InputTag.LMB。"),
            ("AuraInputComponent.h", "自定义 EnhancedInputComponent，批量把 AbilityInputActions 绑定到 Pressed、Released、Held 三类函数。"),
            ("AuraPlayerController.h/.cpp", "输入系统核心调度者：鼠标检测、左键两用、短按寻路、长按移动、自动跑、把技能输入转发给 ASC。"),
            ("AuraAbilitySystemComponent.h/.cpp", "技能输入核心：把 StartupInputTag 塞入 AbilitySpec.DynamicAbilityTags；输入发生时遍历技能并激活匹配者。"),
            ("AuraGameplayAbility.h", "所有 Aura 技能的基类，新增 StartupInputTag，让每个技能蓝图声明自己响应哪个输入 Tag。"),
            ("AuraProjectileSpell.h/.cpp", "投射物技能 Ability：激活时在服务器获取 CombatSocketLocation 并生成 Projectile。"),
            ("AuraProjectile.h/.cpp", "投射物 Actor：Sphere 碰撞体、ProjectileMovement 飞行组件、Overlap 命中入口。"),
            ("CombatInterface.h/.cpp", "提供 GetCombatSocketLocation 接口，让 Ability 不关心 AvatarActor 是玩家还是敌人。"),
            ("AuraCharacterBase.h/.cpp", "实现 CombatInterface，使用 WeaponTipSocketName 从 Weapon 网格体插槽取生成位置。"),
            ("Aura.Build.cs", "模块依赖：EnhancedInput、GameplayAbilities、GameplayTags、GameplayTasks、NavigationSystem 等。"),
        ],
    )

    doc.add_heading("3. InputTag 与 Enhanced Input", level=1)
    doc.add_heading("3.1 InputTag 注册", level=2)
    doc.add_paragraph(
        "FAuraGameplayTags 里新增了一组输入标签变量，用来把具体按键抽象成 GameplayTag。"
    )
    add_code(
        doc,
        "FGameplayTag InputTag_LMB;\n"
        "FGameplayTag InputTag_RMB;\n"
        "FGameplayTag InputTag_1;\n"
        "FGameplayTag InputTag_2;\n"
        "FGameplayTag InputTag_3;\n"
        "FGameplayTag InputTag_4;",
    )
    add_code(
        doc,
        "GameplayTags.InputTag_LMB = UGameplayTagsManager::Get().AddNativeGameplayTag(\n"
        "    FName(\"InputTag.LMB\"),\n"
        "    FString(\"Input Tag for Left mouse button\")\n"
        ");",
    )
    add_bullets(
        doc,
        [
            "InputTag_LMB 不是鼠标左键本身，而是“左键这个输入语义”的标签。",
            "后续技能蓝图的 StartupInputTag 也要填同一个 Tag，ASC 才能把输入和技能配上。",
            "如果编译器不认识 FAuraGameplayTags，要在使用文件 include \"AuraGameplayTags.h\"。",
        ],
    )

    doc.add_heading("3.2 AuraInputConfig：按键动作到 Tag", level=2)
    add_code(
        doc,
        "USTRUCT(BlueprintType)\n"
        "struct FAuraInputAction\n"
        "{\n"
        "    UPROPERTY(EditDefaultsOnly)\n"
        "    const UInputAction* InputAction = nullptr;\n\n"
        "    UPROPERTY(EditDefaultsOnly)\n"
        "    FGameplayTag InputTag = FGameplayTag();\n"
        "};",
    )
    doc.add_paragraph(
        "蓝图 DataAsset 中配置多行 FAuraInputAction，典型配置是 IA_LMB -> InputTag.LMB，IA_1 -> InputTag.1。这样以后换键只改资产，不改技能逻辑。"
    )

    doc.add_heading("3.3 AuraInputComponent：批量绑定", level=2)
    add_code(
        doc,
        "BindAction(Action.InputAction, ETriggerEvent::Started,   Object, PressedFunc,  Action.InputTag);\n"
        "BindAction(Action.InputAction, ETriggerEvent::Completed, Object, ReleasedFunc, Action.InputTag);\n"
        "BindAction(Action.InputAction, ETriggerEvent::Triggered, Object, HeldFunc,     Action.InputTag);",
    )
    add_bullets(
        doc,
        [
            "Started 对应按下瞬间，进入 AbilityInputTagPressed。",
            "Completed 对应松开，进入 AbilityInputTagReleased。",
            "Triggered 对应按住/持续触发，进入 AbilityInputTagHeld。",
            "Action.InputTag 会作为额外参数传给 PlayerController 的回调函数。",
        ],
    )

    doc.add_heading("4. PlayerController：左键两用与寻路自动跑", level=1)
    doc.add_heading("4.1 SetupInputComponent 绑定入口", level=2)
    add_code(
        doc,
        "UAuraInputComponent* AuraInputComponent = CastChecked<UAuraInputComponent>(InputComponent);\n"
        "AuraInputComponent->BindAction(MoveAction, ETriggerEvent::Triggered, this, &AAuraPlayerController::Move);\n"
        "AuraInputComponent->BindAbilityActions(\n"
        "    InputConfig,\n"
        "    this,\n"
        "    &ThisClass::AbilityInputTagPressed,\n"
        "    &ThisClass::AbilityInputTagReleased,\n"
        "    &ThisClass::AbilityInputTagHeld\n"
        ");",
    )
    doc.add_paragraph(
        "移动输入仍然直接走 MoveAction；技能输入统一走 BindAbilityActions，再由 InputTag 决定后续行为。"
    )

    doc.add_heading("4.2 CursorTrace：鼠标下是否是敌人", level=2)
    add_code(
        doc,
        "GetHitResultUnderCursor(ECC_Visibility, false, CursorHit);\n"
        "LastActor = ThisActor;\n"
        "ThisActor = Cast<IEnemyInterface>(CursorHit.GetActor());\n"
        "if (LastActor) LastActor->UnHighlightActor();\n"
        "if (ThisActor) ThisActor->HighlightActor();",
    )
    doc.add_paragraph(
        "ThisActor 是否为空会影响左键行为：鼠标下有敌人时左键被视为技能输入；没有敌人时左键被视为移动输入。"
    )

    doc.add_heading("4.3 AbilityInputTagPressed：记录这次左键是在点谁", level=2)
    add_code(
        doc,
        "if (InputTag.MatchesTagExact(FAuraGameplayTags::Get().InputTag_LMB))\n"
        "{\n"
        "    bTargeting = ThisActor ? true : false;\n"
        "    bAutoRunning = false;\n"
        "}",
    )
    add_bullets(
        doc,
        [
            "bTargeting=true：按下左键时鼠标下有敌人，这次左键后续走技能逻辑。",
            "bTargeting=false：按下左键时没有敌人，这次左键后续走移动/寻路逻辑。",
            "bAutoRunning=false：一旦重新按下左键，就取消之前的自动跑。",
        ],
    )

    doc.add_heading("4.4 AbilityInputTagHeld：非左键直接走技能，左键分流", level=2)
    add_code(
        doc,
        "if (!InputTag.MatchesTagExact(FAuraGameplayTags::Get().InputTag_LMB))\n"
        "{\n"
        "    if (GetASC()) GetASC()->AbilityInputTagHeld(InputTag);\n"
        "    return;\n"
        "}\n\n"
        "if (bTargeting)\n"
        "{\n"
        "    if (GetASC()) GetASC()->AbilityInputTagHeld(InputTag);\n"
        "}\n"
        "else\n"
        "{\n"
        "    FollowTime += GetWorld()->GetDeltaSeconds();\n"
        "    if (CursorHit.bBlockingHit) CachedDestination = CursorHit.ImpactPoint;\n"
        "    AddMovementInput toward CachedDestination;\n"
        "}",
    )
    add_callout(
        doc,
        "理解重点",
        "非 LMB 的技能键没有“点地移动”语义，直接转给 ASC；LMB 则要看 bTargeting，如果点敌人就转 ASC，如果点地面就持续向鼠标位置移动。",
    )

    doc.add_heading("4.5 AbilityInputTagReleased：短按寻路，目标状态释放技能", level=2)
    add_code(
        doc,
        "if (!InputTag.MatchesTagExact(FAuraGameplayTags::Get().InputTag_LMB))\n"
        "{\n"
        "    if (GetASC()) GetASC()->AbilityInputTagReleased(InputTag);\n"
        "    return;\n"
        "}\n\n"
        "if (bTargeting)\n"
        "{\n"
        "    if (GetASC()) GetASC()->AbilityInputTagReleased(InputTag);\n"
        "}\n"
        "else if (FollowTime <= ShortPressThreshold && ControlledPawn)\n"
        "{\n"
        "    FindPathToLocationSynchronously(...);\n"
        "    Spline->ClearSplinePoints();\n"
        "    AddSplinePoint for each NavPath point;\n"
        "    CachedDestination = last path point;\n"
        "    bAutoRunning = true;\n"
        "}",
    )
    add_bullets(
        doc,
        [
            "非 LMB 松开：通知 ASC 输入释放。",
            "LMB 且 bTargeting=true：说明这次是点敌人，通知 ASC 释放输入。",
            "LMB 且 bTargeting=false：说明这次是点地面。若按下时间小于 ShortPressThreshold，就算短按，计算导航路径并开启自动跑。",
            "NavPath->PathPoints.Num() 必须大于 0 才能取最后一个点，否则会出现数组 -1 越界断言。",
        ],
    )

    doc.add_heading("4.6 AutoRun：沿 Spline 自动移动", level=2)
    add_code(
        doc,
        "if (!bAutoRunning) return;\n"
        "const FVector LocationOnSpline = Spline->FindLocationClosestToWorldLocation(\n"
        "    ControllerPawn->GetActorLocation(), ESplineCoordinateSpace::World);\n"
        "const FVector Direction = Spline->FindDirectionClosestToWorldLocation(\n"
        "    LocationOnSpline, ESplineCoordinateSpace::World);\n"
        "ControllerPawn->AddMovementInput(Direction);\n\n"
        "if ((LocationOnSpline - CachedDestination).Length() <= AutoRunAcceptanceRadius)\n"
        "{\n"
        "    bAutoRunning = false;\n"
        "}",
    )
    add_bullets(
        doc,
        [
            "Spline 记录导航系统算出来的路径点。",
            "FindLocationClosestToWorldLocation 用角色世界坐标找 Spline 上最近点。",
            "FindDirectionClosestToWorldLocation 取当前位置附近的路径方向。",
            "AddMovementInput 让角色沿路径方向移动。",
            "接近 CachedDestination 后关闭 bAutoRunning。",
        ],
    )

    doc.add_heading("5. ASC：用 InputTag 找技能", level=1)
    doc.add_heading("5.1 技能被授予时塞入 DynamicAbilityTags", level=2)
    add_code(
        doc,
        "FGameplayAbilitySpec AbilitySpec = FGameplayAbilitySpec(AbilityClass, 1);\n"
        "if (UAuraGameplayAbility* AuraAbility = Cast<UAuraGameplayAbility>(AbilitySpec.Ability))\n"
        "{\n"
        "    AbilitySpec.DynamicAbilityTags.AddTag(AuraAbility->StartupInputTag);\n"
        "    GiveAbility(AbilitySpec);\n"
        "}",
    )
    doc.add_paragraph(
        "每个技能蓝图继承 UAuraGameplayAbility，并配置 StartupInputTag。GiveAbility 之前把这个 Tag 塞进 AbilitySpec.DynamicAbilityTags，之后 ASC 就能根据输入 Tag 找到这个技能。"
    )

    doc.add_heading("5.2 Held 时激活技能", level=2)
    add_code(
        doc,
        "for (FGameplayAbilitySpec& AbilitySpec : GetActivatableAbilities())\n"
        "{\n"
        "    if (AbilitySpec.DynamicAbilityTags.HasTagExact(InputTag))\n"
        "    {\n"
        "        AbilitySpecInputPressed(AbilitySpec);\n"
        "        if (!AbilitySpec.IsActive())\n"
        "        {\n"
        "            TryActivateAbility(AbilitySpec.Handle);\n"
        "        }\n"
        "    }\n"
        "}",
    )
    add_bullets(
        doc,
        [
            "GetActivatableAbilities 返回 ASC 当前拥有的所有技能 Spec。",
            "DynamicAbilityTags.HasTagExact(InputTag) 是输入和技能配对的核心判断。",
            "AbilitySpecInputPressed 告诉 GAS：这个技能对应的输入正在按下。",
            "TryActivateAbility 使用 AbilitySpec.Handle 激活对应技能。",
        ],
    )

    doc.add_heading("5.3 Released 时通知技能释放", level=2)
    add_code(
        doc,
        "if (AbilitySpec.DynamicAbilityTags.HasTagExact(InputTag))\n"
        "{\n"
        "    AbilitySpecInputReleased(AbilitySpec);\n"
        "}",
    )
    doc.add_paragraph(
        "Released 不负责激活技能，而是告诉 GAS 输入松开。蓄力、引导、按住释放类技能后面会依赖这个信号。"
    )

    doc.add_heading("6. ProjectileSpell：技能激活后生成投射物", level=1)
    doc.add_heading("6.1 Ability 蓝图的 ProjectileClass", level=2)
    add_code(
        doc,
        "UPROPERTY(EditAnywhere, BlueprintReadOnly)\n"
        "TSubclassOf<AAuraProjectile> ProjectileClass;",
    )
    add_bullets(
        doc,
        [
            "ProjectileClass 存的是投射物 Actor 类，不是 Ability 类。",
            "正确类型是 TSubclassOf<AAuraProjectile>。",
            "如果写成 AAuraProjectileSpell，UHT 会找不到这个 A 类，因为技能类实际是 UAuraProjectileSpell。",
        ],
    )

    doc.add_heading("6.2 ActivateAbility 只在服务器生成 Projectile", level=2)
    add_code(
        doc,
        "const bool bIsServer = HasAuthority(&ActivationInfo);\n"
        "if (!bIsServer) return;\n\n"
        "ICombatInterface* CombatInterface = Cast<ICombatInterface>(GetAvatarActorFromActorInfo());\n"
        "if (CombatInterface)\n"
        "{\n"
        "    const FVector SocketLocation = CombatInterface->GetCombatSocketLocation();\n"
        "    FTransform SpawnTransform;\n"
        "    SpawnTransform.SetLocation(SocketLocation);\n"
        "    AAuraProjectile* Projectile = GetWorld()->SpawnActorDeferred<AAuraProjectile>(...);\n"
        "    Projectile->FinishSpawning(SpawnTransform);\n"
        "}",
    )
    add_bullets(
        doc,
        [
            "HasAuthority(&ActivationInfo) 防止客户端也生成投射物，避免重复生成。",
            "GetAvatarActorFromActorInfo 返回当前 Ability 的实际表现 Actor，通常是角色。",
            "Cast<ICombatInterface> 让 Ability 不关心角色具体类型，只通过接口拿战斗插槽位置。",
            "SpawnActorDeferred 允许先创建对象、配置属性，再 FinishSpawning 完成生成。",
        ],
    )

    doc.add_heading("6.3 CombatInterface 与武器插槽", level=2)
    add_code(
        doc,
        "virtual FVector GetCombatSocketLocation() override;\n\n"
        "FVector AAuraCharacterBase::GetCombatSocketLocation()\n"
        "{\n"
        "    return Weapon->GetSocketLocation(WeaponTipSocketName);\n"
        "}",
    )
    add_callout(
        doc,
        "接口意义",
        "ProjectileSpell 只知道 AvatarActor 实现了 ICombatInterface，不需要判断它是玩家还是敌人。真正执行的是具体角色类重写的 GetCombatSocketLocation。",
    )

    doc.add_heading("7. AuraProjectile：投射物 Actor", level=1)
    doc.add_heading("7.1 组件结构", level=2)
    add_code(
        doc,
        "Sphere = CreateDefaultSubobject<USphereComponent>(\"Sphere\");\n"
        "SetRootComponent(Sphere);\n"
        "Sphere->SetCollisionEnabled(ECollisionEnabled::QueryOnly);\n"
        "Sphere->SetCollisionResponseToAllChannels(ECR_Ignore);\n"
        "Sphere->SetCollisionResponseToChannel(ECC_WorldDynamic, ECR_Overlap);\n"
        "Sphere->SetCollisionResponseToChannel(ECC_WorldStatic, ECR_Overlap);\n"
        "Sphere->SetCollisionResponseToChannel(ECC_Pawn, ECR_Overlap);\n\n"
        "ProjectileMovement = CreateDefaultSubobject<UProjectileMovementComponent>(\"ProjectileMovement\");\n"
        "ProjectileMovement->InitialSpeed = 550.f;\n"
        "ProjectileMovement->MaxSpeed = 550.f;\n"
        "ProjectileMovement->ProjectileGravityScale = 0.f;",
    )
    add_bullets(
        doc,
        [
            "Sphere 是根组件，也是投射物的碰撞体。",
            "QueryOnly 表示只做查询/重叠，不做物理模拟碰撞。",
            "ProjectileMovement 负责飞行速度和重力行为。",
            "ProjectileGravityScale=0 表示投射物不受重力下坠影响。",
        ],
    )

    doc.add_heading("7.2 Overlap 绑定", level=2)
    add_code(
        doc,
        "Sphere->OnComponentBeginOverlap.AddDynamic(\n"
        "    this,\n"
        "    &AAuraProjectile::OnSphereOverlap\n"
        ");",
    )
    doc.add_paragraph(
        "这句把 Sphere 的开始重叠事件绑定到 OnSphereOverlap。之后投射物碰到 Pawn、WorldStatic 或 WorldDynamic 时，会自动进入这个函数。"
    )

    doc.add_heading("8. 两条完整流程", level=1)
    doc.add_heading("8.1 短按左键点地自动跑", level=2)
    add_numbered(
        doc,
        [
            "玩家按下 LMB，AbilityInputTagPressed 判断当前 ThisActor 是否为空。",
            "如果没有敌人，bTargeting=false，并停止旧的 bAutoRunning。",
            "玩家按住期间，AbilityInputTagHeld 不走 ASC，而是持续更新 CachedDestination 并 AddMovementInput。",
            "玩家很快松开 LMB，AbilityInputTagReleased 判断 FollowTime <= ShortPressThreshold。",
            "通过 UNavigationSystemV1::FindPathToLocationSynchronously 计算从角色到 CachedDestination 的导航路径。",
            "把 NavPath->PathPoints 写入 Spline。",
            "设置 CachedDestination 为路径最后一点，bAutoRunning=true。",
            "PlayerTick 每帧调用 AutoRun，角色沿 Spline 方向移动。",
            "距离终点小于 AutoRunAcceptanceRadius 时，bAutoRunning=false。",
        ],
    )

    doc.add_heading("8.2 左键点敌人生成投射物技能", level=2)
    add_numbered(
        doc,
        [
            "CursorTrace 检测鼠标下 Actor，如果实现 IEnemyInterface，ThisActor 非空并高亮。",
            "玩家按下 LMB，bTargeting=true。",
            "玩家按住 LMB，AbilityInputTagHeld 把 InputTag.LMB 转交给 AuraASC。",
            "AuraASC 遍历 GetActivatableAbilities，找到 DynamicAbilityTags 中含 InputTag.LMB 的 AbilitySpec。",
            "ASC 调用 AbilitySpecInputPressed，并 TryActivateAbility。",
            "UAuraProjectileSpell::ActivateAbility 执行。",
            "服务器端通过 ICombatInterface 获取 WeaponTipSocketName 的世界位置。",
            "SpawnActorDeferred 生成 AAuraProjectile，FinishSpawning 完成生成。",
            "ProjectileMovement 让投射物飞行，Sphere 重叠事件等待后续命中逻辑。",
        ],
    )

    doc.add_heading("9. 常见错误与排查", level=1)
    add_three_col_table(
        doc,
        ("现象", "常见原因", "处理方式"),
        [
            ("FAuraGameplayTags 后面 :: 报错", "当前 cpp 没包含 AuraGameplayTags.h", "加 #include \"AuraGameplayTags.h\"。"),
            ("FGameplayTag class/struct 冲突", "把 FGameplayTag 前向声明成 class", "删掉错误前向声明，include GameplayTagContainer.h。"),
            ("FindPathToLocationSynchronously 链接错误 LNK2019", "Build.cs 没链接 NavigationSystem", "Aura.Build.cs 加 NavigationSystem 模块依赖。"),
            ("PathPoints[-1] 断言", "NavPath 存在但 PathPoints.Num()==0", "访问 Last 前判断 Num()>0，最好 Num()>1。"),
            ("USplineComponent 未定义", "只前向声明但 cpp 调用了成员函数", "cpp include Components/SplineComponent.h。"),
            ("ProjectileMovement 未定义", "只前向声明但访问 InitialSpeed 等成员", "cpp include GameFramework/ProjectileMovementComponent.h。"),
            ("Unable to find class AAuraProjectileSpell", "ProjectileClass 写成了不存在的 A 类", "改成 TSubclassOf<AAuraProjectile>。"),
            ("Projectile 运行时不生成", "ProjectileClass 蓝图未配置，或客户端执行被 HasAuthority return", "确认技能蓝图 ProjectileClass 已指定；服务器端激活路径正常。"),
            ("左键点地却放技能", "CursorTrace 中 ThisActor 没及时清空或检测通道不对", "检查 ECC_Visibility 响应和 ThisActor 更新逻辑。"),
            ("左键点敌人不放技能", "技能 StartupInputTag 与 InputConfig 的 InputTag.LMB 不一致", "确认 GA 蓝图和 InputConfig 使用完全相同 Tag。"),
        ],
    )

    doc.add_heading("10. 和已有文档的关系", level=1)
    add_kv_table(
        doc,
        [
            ("HUD / WidgetController 文档", "讲 UI Controller 如何创建、绑定、广播。新链路中 PlayerController 扮演输入调度者，ASC 扮演技能调度者。"),
            ("AttributeMenu 文档", "讲 Tag 如何把属性和 UI 行配对。新链路中 Tag 把 InputAction 和 GameplayAbility 配对。"),
            ("GameplayEffect 文档", "讲 GE、Spec、Context、应用和回调。Projectile 命中后后续大概率会把伤害 GE 应用到目标，这会接回 GameplayEffect 链路。"),
            ("本篇文档", "聚焦玩家输入到技能激活，再到投射物 Actor 生成与自动寻路移动。"),
        ],
    )

    doc.add_heading("11. 推荐调试顺序", level=1)
    add_numbered(
        doc,
        [
            "先确认 InputConfig 资产里每个 InputAction 都配置了有效 InputTag。",
            "在 AbilityInputTagPressed/Held/Released 打印 InputTag，确认按键能进 PlayerController。",
            "在 ASC 的 AbilityInputTagHeld 打印每个 AbilitySpec.DynamicAbilityTags，确认技能被 GiveAbility 且带着正确 Tag。",
            "确认 GA 蓝图的 StartupInputTag 与 InputConfig 里的 InputTag 完全一致。",
            "ProjectileSpell 中打印 bIsServer，确认投射物只在服务器生成。",
            "打印 CombatInterface->GetCombatSocketLocation，确认武器插槽位置正确。",
            "Projectile BeginPlay 中打印，确认投射物 Actor 已生成。",
            "OnSphereOverlap 中打印 OtherActor，确认碰撞通道和重叠事件正常。",
            "自动跑异常时，打印 PathPoints.Num、Spline->GetNumberOfSplinePoints、CachedDestination、DistanceToDestination。",
        ],
    )

    doc.add_heading("12. 后续扩展点", level=1)
    add_bullets(
        doc,
        [
            "给 AutoRun 加 Spline 点数保护：bAutoRunning 开启前和 AutoRun 内都确认 Spline 点数大于 0。",
            "ProjectileSpell 中根据鼠标命中点计算旋转，让投射物朝目标方向飞，而不是只设置生成位置。",
            "在 Projectile 上保存 DamageEffectSpecHandle 或 DamageParams，Overlap 时给目标应用伤害 GE。",
            "在 AuraASC 中改用 UE5.5 推荐的 GetDynamicSpecSourceTags，替代过期的 DynamicAbilityTags 直接访问。",
            "为 LMB 区分点击移动、点击攻击、长按移动、技能引导等更多输入状态。",
        ],
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Aura InputTag AutoRun Projectile Technical Guide")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string("666666")

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
