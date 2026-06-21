from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT = "Docs/Aura_GameplayEffect_Technical_Guide_v2_Naming.docx"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Microsoft YaHei"
    r.font.size = Pt(9.5)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, header, True)
        shade(cell, "EAF2F8")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if widths:
            cell.width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths:
                cells[i].width = Inches(widths[i])
    doc.add_paragraph()


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(35, 35, 35)


def add_note(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    shade(cell, "F6F8FA")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Microsoft YaHei"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string("263238")
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.85)
section.bottom_margin = Inches(0.85)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)

styles = doc.styles
styles["Normal"].font.name = "Microsoft YaHei"
styles["Normal"].font.size = Pt(10.5)
styles["Normal"].paragraph_format.space_after = Pt(6)
styles["Normal"].paragraph_format.line_spacing = 1.2

for name, size, color in [
    ("Heading 1", 16, "1F4E79"),
    ("Heading 2", 13, "2F6F9F"),
    ("Heading 3", 11.5, "234B5E"),
]:
    style = styles[name]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(12)
    style.paragraph_format.space_after = Pt(5)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Aura GameplayEffect 技术文档 v2")
run.bold = True
run.font.name = "Microsoft YaHei"
run.font.size = Pt(20)
run.font.color.rgb = RGBColor.from_string("0B2545")

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run("基于当前代码：EffectActor、AttributeSet、ASC 标签广播、UI 消息行与命名建议")
subtitle_run.font.name = "Microsoft YaHei"
subtitle_run.font.size = Pt(10)
subtitle_run.font.color.rgb = RGBColor.from_string("586069")

add_note(
    doc,
    "一句话总览：GameplayEffect 是效果配方，ASC 负责应用和管理，AttributeSet 保存属性并做收尾，EffectActor 负责触发，WidgetController 负责把 GE 标签转成 UI 消息。"
)

doc.add_heading("1. 当前系统总览", level=1)
doc.add_paragraph(
    "你现在的 GameplayEffect 流程已经从“一个 Actor 只配置一个 GE”进化到“按 Instant / Duration / Infinite 分类配置多个 GE”。"
    "这是一个更适合扩展的方向：同一个药水、区域、陷阱或 Buff Actor 可以同时施加多种效果。"
)
add_code(
    doc,
    "EffectActor 蓝图重叠\n"
    "    -> AAuraEffectActor::OnOverlap(TargetActor)\n"
    "        -> 根据 ApplicationPolicy 遍历 Instant / Duration / Infinite GE 数组\n"
    "        -> ApplyEffectToTarget(TargetActor, EffectClass)\n"
    "            -> 取得 TargetASC\n"
    "            -> MakeEffectContext\n"
    "            -> MakeOutgoingSpec\n"
    "            -> ApplyGameplayEffectSpecToSelf\n"
    "            -> 如果是 Infinite 且需要 EndOverlap 移除，保存 ActiveEffectHandle\n\n"
    "EffectActor 蓝图结束重叠\n"
    "    -> AAuraEffectActor::OnEndOverlap(TargetActor)\n"
    "        -> 可按策略应用 EndOverlap GE\n"
    "        -> 查找 TargetASC 对应的 Infinite Handles\n"
    "        -> RemoveActiveGameplayEffect\n"
    "        -> 从 ActiveEffectHandles 中删除记录"
)

doc.add_heading("2. 文件职责", level=1)
add_table(
    doc,
    ["文件", "主要内容", "负责什么"],
    [
        [
            "Source/Aura/Public/Actor/AuraEffectActor.h",
            "AAuraEffectActor、EEffectApplicationPolicy、EEffectRemovalPolicy",
            "声明效果应用策略、移除策略、GE 数组、ActiveEffectHandles 和 ActorLevel。",
        ],
        [
            "Source/Aura/Private/Actor/AuraEffectActor.cpp",
            "ApplyEffectToTarget / OnOverlap / OnEndOverlap",
            "执行 GE 应用；保存 Infinite Handle；离开范围时移除对应 Infinite 效果。",
        ],
        [
            "Source/Aura/Public/AbilitySystem/AuraAttributeSet.h",
            "Health / MaxHealth / Mana / MaxMana、FEffectProperties",
            "声明属性、属性访问宏、复制回调，以及 GE 来源/目标信息结构。",
        ],
        [
            "Source/Aura/Private/AbilitySystem/AuraAttributeSet.cpp",
            "PreAttributeChange / PostGameplayEffectExecute / SetEffectProperties",
            "限制属性范围，GE 执行后收尾，并整理 Source / Target 信息。",
        ],
        [
            "Source/Aura/Public/AbilitySystem/AuraAbilitySystemComponent.h",
            "FEffectAssetTags、EffectAssetTags",
            "声明 GE AssetTags 广播，让 UI 或其他系统监听 GE 标签。",
        ],
        [
            "Source/Aura/Private/AbilitySystem/AuraAbilitySystemComponent.cpp",
            "AbilityActorInfoSet / EffectApplied",
            "绑定 GE 应用回调，从 EffectSpec 取 AssetTags 并广播。",
        ],
        [
            "Source/Aura/Public/UI/WidgetController/OverlayWidgetController.h",
            "FUIWidgetRow、MessageWidgetRowDelegate",
            "声明 UI 消息表行结构和消息广播。",
        ],
        [
            "Source/Aura/Private/UI/WidgetController/OverlayWidgetController.cpp",
            "BindCallbacksToDependencies",
            "监听 ASC 的 EffectAssetTags，根据标签查 DataTable，再广播给 WBP 创建消息 Widget。",
        ],
    ],
    [2.4, 2.0, 2.8],
)

doc.add_heading("3. AuraEffectActor：效果触发器", level=1)
doc.add_paragraph("文件：Source/Aura/Public/Actor/AuraEffectActor.h")

doc.add_heading("3.1 应用策略", level=2)
add_code(
    doc,
    "UENUM(BlueprintType)\n"
    "enum class EEffectApplicationPolicy : uint8\n"
    "{\n"
    "    ApplyOnOverlap,\n"
    "    ApplyOnEndOverlap,\n"
    "    DoNotApply\n"
    "};"
)
add_bullets(
    doc,
    [
        "ApplyOnOverlap：进入碰撞范围时应用该类效果。",
        "ApplyOnEndOverlap：离开碰撞范围时应用该类效果。",
        "DoNotApply：这类效果不自动应用。",
        "BlueprintType 的 enum class 必须指定 : uint8，否则 UHT 会报 Invalid BlueprintType enum base。",
    ],
)

doc.add_heading("3.2 移除策略", level=2)
add_code(
    doc,
    "UENUM(BlueprintType)\n"
    "enum class EEffectRemovalPolicy : uint8\n"
    "{\n"
    "    RemoveOnEndOverlap,\n"
    "    DoNotRemove\n"
    "};"
)
add_bullets(
    doc,
    [
        "RemoveOnEndOverlap：离开范围时移除之前保存的 Infinite Effects。",
        "DoNotRemove：离开范围后不移除，Infinite 效果会继续留在目标 ASC 上。",
    ],
)

doc.add_heading("3.3 GE 数组", level=2)
add_code(
    doc,
    "UPROPERTY(EditAnywhere, BlueprintReadOnly, Category = \"Applied Effects\")\n"
    "TArray<TSubclassOf<UGameplayEffect>> InstantGameplayEffectClasses;\n\n"
    "UPROPERTY(EditAnywhere, BlueprintReadOnly, Category = \"Applied Effects\")\n"
    "TArray<TSubclassOf<UGameplayEffect>> DurationGameplayEffectClasses;\n\n"
    "UPROPERTY(EditAnywhere, BlueprintReadOnly, Category = \"Applied Effects\")\n"
    "TArray<TSubclassOf<UGameplayEffect>> InfiniteGameplayEffectClasses;"
)
add_bullets(
    doc,
    [
        "数组版允许一个 EffectActor 同时应用多个 GE。",
        "蓝图里只需要在 Details 面板配置数组，不需要每次手动给 ApplyEffectToTarget 传 Class。",
        "C++ 的 OnOverlap / OnEndOverlap 负责统一遍历和应用。",
    ],
)
add_note(
    doc,
    "当前头文件里还残留了 InstantGameplayEffectClass / DurationGameplayEffectClass / InfiniteGameplayEffectClass 这三个单个字段，而且没有 UPROPERTY。既然你已经改用数组，建议后续删掉旧字段，避免自己在蓝图和 C++ 里看混。"
)

doc.add_heading("4. ApplyEffectToTarget：真正应用 GE 的地方", level=1)
doc.add_paragraph("文件：Source/Aura/Private/Actor/AuraEffectActor.cpp")
doc.add_paragraph("函数：AAuraEffectActor::ApplyEffectToTarget(AActor* TargetActor, TSubclassOf<UGameplayEffect> GameplayEffectClass)")
add_code(
    doc,
    "UAbilitySystemComponent* TargetASC = UAbilitySystemBlueprintLibrary::GetAbilitySystemComponent(TargetActor);\n"
    "if (TargetASC == nullptr) return;\n"
    "if (!GameplayEffectClass) return;\n\n"
    "FGameplayEffectContextHandle EffectContextHandle = TargetASC->MakeEffectContext();\n"
    "EffectContextHandle.AddSourceObject(this);\n\n"
    "FGameplayEffectSpecHandle EffectSpecHandle = TargetASC->MakeOutgoingSpec(GameplayEffectClass, ActorLevel, EffectContextHandle);\n"
    "FActiveGameplayEffectHandle ActiveEffectHandle = TargetASC->ApplyGameplayEffectSpecToSelf(*EffectSpecHandle.Data.Get());"
)
add_table(
    doc,
    ["步骤", "含义"],
    [
        ["GetAbilitySystemComponent", "从目标 Actor 上取得 ASC。没有 ASC 就无法应用 GE。"],
        ["MakeEffectContext", "创建本次效果的上下文，用来记录来源、Instigator、SourceObject 等背景信息。"],
        ["AddSourceObject(this)", "把当前 AuraEffectActor 记录成本次效果来源对象。"],
        ["MakeOutgoingSpec", "用 GE Class、ActorLevel 和 Context 生成本次具体效果实例。"],
        ["ApplyGameplayEffectSpecToSelf", "把这份 Spec 应用到目标 ASC 自己身上，并返回 ActiveEffectHandle。"],
    ],
    [2.1, 4.8],
)

doc.add_heading("5. Context / Spec / Handle", level=1)
add_table(
    doc,
    ["概念", "直观理解", "作用"],
    [
        ["UGameplayEffect / GameplayEffectClass", "配方", "定义改什么属性、怎么改、持续多久、是否周期、是否堆叠。"],
        ["FGameplayEffectContextHandle", "现场记录", "记录来源对象、Instigator、HitResult 等上下文信息。"],
        ["FGameplayEffectSpecHandle", "本次执行单据", "由 GE 配方 + 等级 + Context 生成的具体运行时效果。"],
        ["FActiveGameplayEffectHandle", "已激活效果的编号", "应用成功后用于之后查找或移除这个 Active GE。"],
        ["UAbilitySystemComponent", "账本和执行器", "保存 ActiveGameplayEffects，并负责应用、复制、移除 GE。"],
    ],
    [1.8, 1.7, 3.4],
)
add_note(
    doc,
    "FActiveGameplayEffectHandle 不是效果本体。真正的 Active GameplayEffect 数据在 ASC 内部，Handle 只是之后找到它的钥匙。"
)

doc.add_heading("6. Infinite GE 的保存与移除", level=1)

doc.add_heading("6.1 保存 Handle", level=2)
add_code(
    doc,
    "bool bIsInfinite = EffectSpecHandle.Data.Get()->Def.Get()->DurationPolicy == EGameplayEffectDurationType::Infinite;\n\n"
    "if (bIsInfinite && InfiniteEffectRemovalPolicy == EEffectRemovalPolicy::RemoveOnEndOverlap)\n"
    "{\n"
    "    ActiveEffectHandles.Add(ActiveEffectHandle, TargetASC);\n"
    "}"
)
add_bullets(
    doc,
    [
        "只有 Infinite GE 需要手动移除，所以只保存 Infinite 的 Handle。",
        "ActiveEffectHandles 的 Key 是 FActiveGameplayEffectHandle，Value 是拥有该效果的 ASC。",
        "保存 ASC 是因为移除时必须对那个 ASC 调用 RemoveActiveGameplayEffect。",
    ],
)

doc.add_heading("6.2 EndOverlap 移除", level=2)
add_code(
    doc,
    "TArray<FActiveGameplayEffectHandle> HandlesToRemove;\n"
    "for (TTuple<FActiveGameplayEffectHandle, UAbilitySystemComponent*> HandlePair : ActiveEffectHandles)\n"
    "{\n"
    "    if (TargetASC == HandlePair.Value)\n"
    "    {\n"
    "        TargetASC->RemoveActiveGameplayEffect(HandlePair.Key, 1);\n"
    "        HandlesToRemove.Add(HandlePair.Key);\n"
    "    }\n"
    "}\n\n"
    "for (FActiveGameplayEffectHandle& Handle : HandlesToRemove)\n"
    "{\n"
    "    ActiveEffectHandles.FindAndRemoveChecked(Handle);\n"
    "}"
)
add_bullets(
    doc,
    [
        "先遍历 Map，找到属于当前 TargetASC 的所有 Handles。",
        "调用 RemoveActiveGameplayEffect 移除目标身上的 Active GE。",
        "不要遍历 Map 时直接删除 Map 元素；先放进 HandlesToRemove，循环结束后再删。",
        "从 TMap 删除 Key 会同时删除对应 Value 记录，但不会销毁 ASC 对象。",
    ],
)

doc.add_heading("7. AttributeSet：属性与 GE 收尾", level=1)
doc.add_paragraph("文件：Source/Aura/Private/AbilitySystem/AuraAttributeSet.cpp")

doc.add_heading("7.1 PreAttributeChange", level=2)
add_code(
    doc,
    "void UAuraAttributeSet::PreAttributeChange(const FGameplayAttribute& Attribute, float& NewValue)\n"
    "{\n"
    "    Super::PreAttributeChange(Attribute, NewValue);\n"
    "    if (Attribute == GetHealthAttribute())\n"
    "    {\n"
    "        NewValue = FMath::Clamp(NewValue, 0.f, GetMaxHealth());\n"
    "    }\n"
    "    if (Attribute == GetManaAttribute())\n"
    "    {\n"
    "        NewValue = FMath::Clamp(NewValue, 0.f, GetMaxMana());\n"
    "    }\n"
    "}"
)
add_bullets(
    doc,
    [
        "这是 GAS 自动调用的生命周期函数，不需要你手动调用。",
        "Attribute 表示即将变化的是哪个属性。",
        "NewValue 是引用，修改它会影响这次变化要写入的值。",
        "它适合做属性变化前的限制，但对 GE 修改后的最终结果，最好还在 PostGameplayEffectExecute 里兜底。",
    ],
)

doc.add_heading("7.2 PostGameplayEffectExecute", level=2)
add_code(
    doc,
    "void UAuraAttributeSet::PostGameplayEffectExecute(const FGameplayEffectModCallbackData& Data)\n"
    "{\n"
    "    Super::PostGameplayEffectExecute(Data);\n"
    "    FEffectProperties Props;\n"
    "    SetEffectProperties(Data, Props);\n\n"
    "    if (Data.EvaluatedData.Attribute == GetHealthAttribute())\n"
    "    {\n"
    "        SetHealth(FMath::Clamp(GetHealth(), 0.f, GetMaxHealth()));\n"
    "    }\n\n"
    "    if (Data.EvaluatedData.Attribute == GetManaAttribute())\n"
    "    {\n"
    "        SetMana(FMath::Clamp(GetMana(), 0.f, GetMaxMana()));\n"
    "    }\n"
    "}"
)
add_bullets(
    doc,
    [
        "GE 修改属性后，GAS 会调用这个函数。",
        "这里适合做最终 Clamp、死亡判断、经验归属、击杀者记录、浮字等后处理。",
        "你现在已经把 Health 和 Mana 限制在 0 到 Max 之间，这是防止药水加过头的关键。",
    ],
)

doc.add_heading("7.3 FEffectProperties", level=2)
add_table(
    doc,
    ["字段", "含义"],
    [
        ["EffectContextHandle", "本次 GE 的上下文。"],
        ["SourceASC", "效果来源方 ASC。"],
        ["SourceAvatarActor", "来源 ASC 的 AvatarActor。"],
        ["SourceController", "来源控制器，玩家可能是 PlayerController，AI 可能是 AIController。"],
        ["SourceCharacter", "来源角色。"],
        ["TargetASC", "被影响目标的 ASC。"],
        ["TargetAvatarActor", "被影响目标的 AvatarActor。"],
        ["TargetController", "目标控制器。"],
        ["TargetCharacter", "目标角色。"],
    ],
    [2.1, 4.8],
)
add_note(
    doc,
    "SetEffectProperties 的目的不是改属性，而是整理信息。以后做伤害来源、击杀归属、经验、浮字、仇恨等逻辑时，Props 会很有用。"
)

doc.add_heading("8. ASC：监听 GE 应用并广播 AssetTags", level=1)
doc.add_paragraph("文件：Source/Aura/Private/AbilitySystem/AuraAbilitySystemComponent.cpp")
add_code(
    doc,
    "void UAuraAbilitySystemComponent::AbilityActorInfoSet()\n"
    "{\n"
    "    OnGameplayEffectAppliedDelegateToSelf.AddUObject(this, &UAuraAbilitySystemComponent::EffectApplied);\n"
    "}\n\n"
    "void UAuraAbilitySystemComponent::EffectApplied(\n"
    "    UAbilitySystemComponent* AbilitySystemComponent,\n"
    "    const FGameplayEffectSpec& EffectSpec,\n"
    "    FActiveGameplayEffectHandle ActiveEffectHandle)\n"
    "{\n"
    "    FGameplayTagContainer TagContainer;\n"
    "    EffectSpec.GetAllAssetTags(TagContainer);\n"
    "    EffectAssetTags.Broadcast(TagContainer);\n"
    "}"
)
add_bullets(
    doc,
    [
        "OnGameplayEffectAppliedDelegateToSelf 是 ASC 自带的 GE 应用事件。",
        "AddUObject(this, &UAuraAbilitySystemComponent::EffectApplied) 表示：有 GE 应用到自己身上时，调用当前 ASC 的 EffectApplied。",
        "EffectSpec.GetAllAssetTags 会取出这个 GE 的 AssetTags。",
        "EffectAssetTags.Broadcast(TagContainer) 把标签广播给外部监听者，比如 OverlayWidgetController。",
    ],
)

doc.add_heading("9. UI 消息：从 GE Tag 到 WBP 消息", level=1)
doc.add_paragraph("文件：Source/Aura/Private/UI/WidgetController/OverlayWidgetController.cpp")
add_code(
    doc,
    "Cast<UAuraAbilitySystemComponent>(AbilitySystemComponent)->EffectAssetTags.AddLambda(\n"
    "    [this](const FGameplayTagContainer& AssetTags)\n"
    "    {\n"
    "        for (const FGameplayTag& Tag : AssetTags)\n"
    "        {\n"
    "            FGameplayTag MessageTag = FGameplayTag::RequestGameplayTag(FName(\"Message\"));\n"
    "            if (Tag.MatchesTag(MessageTag))\n"
    "            {\n"
    "                const FUIWidgetRow* Row = GetDataTableRowByTag<FUIWidgetRow>(MessageWidgetDataTable, Tag);\n"
    "                MessageWidgetRowDelegate.Broadcast(*Row);\n"
    "            }\n"
    "        }\n"
    "    }\n"
    ");"
)
add_bullets(
    doc,
    [
        "OverlayWidgetController 监听 ASC 的 EffectAssetTags。",
        "只处理 Message 开头的标签，例如 Message.HealthPotion、Message.ManaPotion。",
        "用 Tag 的名字去 DataTable 找同名 Row。",
        "找到 FUIWidgetRow 后，通过 MessageWidgetRowDelegate 广播给蓝图。",
        "蓝图收到 Row 后创建对应 MessageWidget，并设置文字和图标。",
    ],
)
add_note(
    doc,
    "这里建议加一个 Row 空指针检查：如果 DataTable 没配、Row 名不匹配或 Tag 不存在，直接 Broadcast(*Row) 会崩。学习阶段可以先记住这个风险。"
)

doc.add_heading("10. 命名建议", level=1)
doc.add_paragraph("命名的目标不是好看，而是让你在 C++、蓝图 Details、DataTable 和日志里一眼知道这个东西属于哪一层。")

add_table(
    doc,
    ["当前名字", "建议名字", "理由"],
    [
        ["InstantGameplayEffectClasses", "InstantGameplayEffectClasses", "可以保留。它明确表示 Instant 类型的 GE Class 数组。"],
        ["DurationGameplayEffectClasses", "DurationGameplayEffectClasses", "可以保留。和 UE 的 Duration Policy 对应。"],
        ["InfiniteGameplayEffectClasses", "InfiniteGameplayEffectClasses", "可以保留。和 Infinite Policy 对应。"],
        ["InstantGameplayEffectClass", "删除或改 DeprecatedInstantGameplayEffectClass", "旧单个字段已不用，容易和数组混淆。最好删。"],
        ["DurationGameplayEffectClass", "删除或改 DeprecatedDurationGameplayEffectClass", "同上。"],
        ["InfiniteGameplayEffectClass", "删除或改 DeprecatedInfiniteGameplayEffectClass", "同上。"],
        ["ActiveEffectHandles", "ActiveEffectHandles", "可以保留。它确实保存 Active GE Handle 到 ASC 的映射。"],
        ["ActorLevel", "EffectLevel 或 ActorLevel", "如果只用于 GE 等级，EffectLevel 更准确；如果教程沿用 ActorLevel，也可以保留。"],
        ["bDestroyOnEffectRemoval", "bDestroyOnEffectApplication 或 bDestroyAfterEffectRemoval", "当前代码没用到。若用于药水吃完销毁，名字要表达销毁时机。"],
        ["FEffectAssetTags", "FEffectAssetTags", "可以保留。它是 GE AssetTags 广播类型。注释里的 AssettTags 拼写建议改为 AssetTags。"],
        ["EffectAssetTags", "EffectAssetTags", "可以保留。它是 ASC 对外广播 GE AssetTags 的委托变量。"],
        ["FUIWidgetRow", "FUIWidgetRow 或 FUIMessageWidgetRow", "如果这个表只服务消息提示，FUIMessageWidgetRow 更具体。"],
        ["MessageWidgetRowDelegate", "MessageWidgetRowDelegate 或 OnMessageWidgetRowReceived", "如果要贴近 UE 事件命名，OnMessageWidgetRowReceived 更像事件。"],
        ["MessageWidgetDataTable", "MessageWidgetDataTable", "可以保留。它清楚表达是消息 Widget 用的数据表。"],
    ],
    [2.2, 2.5, 2.9],
)

doc.add_heading("11. 蓝图资产命名建议", level=1)
add_table(
    doc,
    ["类型", "推荐前缀", "例子"],
    [
        ["GameplayEffect", "GE_", "GE_Potion_HealthInstant、GE_Potion_ManaInstant、GE_Area_FireDamageInfinite"],
        ["EffectActor 蓝图", "BP_", "BP_Potion_Health、BP_Potion_Mana、BP_Area_FireDamage"],
        ["DataTable", "DT_", "DT_MessageWidgetData"],
        ["Widget Blueprint", "WBP_", "WBP_EffectMessage、WBP_HealthGlobe、WBP_ManaGlobe"],
        ["GameplayTag", "层级名", "Message.HealthPotion、Message.ManaPotion、State.Burning"],
    ],
    [1.6, 1.5, 4.5],
)

doc.add_heading("12. 现在这套代码的几个小风险", level=1)
add_numbers(
    doc,
    [
        "AAuraEffectActor.h 里旧的单个 GE Class 字段还留着，但不参与逻辑。后续建议删掉，避免蓝图配置时误以为它们有用。",
        "OverlayWidgetController.cpp 里 MessageWidgetDataTable 查表后没有判断 Row 是否为空。DataTable 没配或 Row 名不等于 Tag 名时会崩。",
        "Cast<UAuraAbilitySystemComponent>(AbilitySystemComponent) 最好确认非空。学习阶段可用 CastChecked，正式代码可用 if 判断。",
        "AddLambda 捕获 [this] 简洁，但生命周期不如 AddUObject / AddWeakLambda 明确。WidgetController 销毁后如果委托仍触发，理论上有风险。",
        "RemoveActiveGameplayEffect(HandlePair.Key, 1) 只移除一层堆叠。如果某个 Infinite GE 允许堆叠，要确认这是不是你想要的。",
    ],
)

doc.add_heading("13. 常见问题排查", level=1)
add_table(
    doc,
    ["现象", "高概率原因", "检查位置"],
    [
        ["进入区域有效，离开不移除", "EndOverlap 没调用 OnEndOverlap，或 ActiveEffectHandles 没保存，或 TargetASC 不匹配。", "EffectActor 蓝图 / AuraEffectActor.cpp"],
        ["ActiveEffectHandles.Add 没执行", "GE 不是 Infinite，或 RemovalPolicy 不是 RemoveOnEndOverlap。", "GE Duration Policy / EffectActor Details"],
        ["血量超过 MaxHealth", "PostGameplayEffectExecute 没对 Health 做最终 Clamp。", "AuraAttributeSet.cpp"],
        ["GE 标签消息不显示", "GE 没配置 AssetTags，或 DataTable Row 名和 Tag 名不一致。", "GE Components / DT_MessageWidgetData"],
        ["蓝图里只有 Source Tags / Target Tags", "UE5.5 的 GE 标签配置已移动到 GameplayEffect Components。", "GE Details -> Components"],
        ["Instant GE 提示 TargetTags 不生效", "Instant GE 不会长期挂在目标身上，Grant Tags to Target Actor 没意义。", "GE Duration Policy"],
        ["UENUM BlueprintType 报错", "enum class 没写 : uint8。", "AuraEffectActor.h"],
        ["FActiveGameplayEffectHandle 做 TMap Key 报 GetTypeHash", "缺 GameplayEffectTypes.h 或相关包含顺序问题。", "AuraEffectActor.h"],
    ],
    [2.0, 3.1, 2.1],
)

doc.add_heading("14. 建议的学习记忆方式", level=1)
add_code(
    doc,
    "GameplayEffect = 配方：定义怎么改属性、持续多久、是否周期、是否堆叠。\n"
    "EffectActor = 触发器：进入/离开范围时应用或移除 GE。\n"
    "AbilitySystemComponent = 执行器和账本：应用 GE，保存 Active GE，负责移除。\n"
    "GameplayEffectSpec = 本次效果实例：GE Class + Level + Context。\n"
    "GameplayEffectContext = 现场信息：来源对象、Instigator、HitResult 等。\n"
    "ActiveGameplayEffectHandle = 钥匙：之后用它移除已应用的效果。\n"
    "AttributeSet = 属性账本：保存 Health/Mana，并在回调里限制或结算属性。\n"
    "WidgetController = UI 中介：监听 ASC 和 AttributeSet，把数据广播给 WBP。"
)

doc.core_properties.title = "Aura GameplayEffect Technical Guide v2"
doc.core_properties.subject = "Aura UE GAS GameplayEffect notes with naming guidance"
doc.core_properties.author = "Codex"
doc.save(OUT)
print(OUT)
