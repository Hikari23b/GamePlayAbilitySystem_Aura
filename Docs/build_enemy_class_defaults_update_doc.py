# -*- coding: utf-8 -*-
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "Docs/Aura_EnemyClass_DefaultAttributes_Update_Guide.docx"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def borders(table, color="D7DEE8", size="4"):
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = tbl_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tbl_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for i, w in enumerate(widths):
            cell = row.cells[i]
            cell.width = Inches(w)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(w * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def run_font(run, size=None, bold=None, color=None, mono=False):
    run.font.name = "Consolas" if mono else "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def style(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        st = doc.styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        st = doc.styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(11)
        st.paragraph_format.left_indent = Inches(0.375)
        st.paragraph_format.first_line_indent = Inches(-0.188)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.25


def title(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("Aura 敌人种类与默认属性初始化更新技术文档")
    run_font(r, 22, True, "0B2545")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("CharacterClassInfo / GameMode defaults / Library initialization / Enemy class defaults")
    run_font(r, 11, False, "555555")

    rows = [
        ("项目路径", r"F:\ueprojiect\Aura"),
        ("本次主题", "敌人的种类与默认属性配置：按 Elementalist / Warrior / Ranger 等类型选择不同 Primary Attributes，并共用 Secondary / Vital 默认属性。"),
        ("承接文档", "Aura_EnemyHealthBar_Projectile_Update_Guide.docx"),
        ("关键资产", r"Content\BluePrints\AbilitySystem\Data\DA_CharacterClassInfo.uasset"),
        ("关键结论", "敌人蓝图只需要配置 CharacterClass 和 Level；真正应用哪些 GameplayEffect，由 GameMode 持有的数据资产和 AuraAbilitySystemLibrary 统一完成。"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    borders(table)
    width(table, [1.7, 4.8])
    for row, (k, v) in zip(table.rows, rows):
        shade(row.cells[0], "E8EEF5")
        row.cells[0].paragraphs[0].add_run(k).bold = True
        row.cells[1].paragraphs[0].add_run(v)


def callout(doc, head, body):
    table = doc.add_table(rows=1, cols=1)
    borders(table, "B8C7D9")
    width(table, [6.5])
    cell = table.cell(0, 0)
    shade(cell, "F4F6F9")
    p = cell.paragraphs[0]
    r = p.add_run(head)
    run_font(r, bold=True, color="1F3A5F")
    p.add_run("  " + body)
    doc.add_paragraph()


def kv(doc, rows, headers=("位置 / 概念", "作用"), widths=(1.95, 4.55)):
    table = doc.add_table(rows=1, cols=2)
    borders(table)
    width(table, list(widths))
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        shade(cell, "E8EEF5")
        for r in cell.paragraphs[0].runs:
            r.bold = True
    for a, b in rows:
        cells = table.add_row().cells
        cells[0].text = a
        cells[1].text = b
        for cell in cells:
            margins(cell)
    doc.add_paragraph()


def tri(doc, headers, rows, widths=(1.45, 2.35, 2.7)):
    table = doc.add_table(rows=1, cols=3)
    borders(table)
    width(table, list(widths))
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        shade(cell, "E8EEF5")
        for r in cell.paragraphs[0].runs:
            r.bold = True
    for row_data in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row_data):
            cells[i].text = text
            margins(cells[i])
    doc.add_paragraph()


def flow(doc, heading, rows):
    doc.add_heading(heading, level=3)
    table = doc.add_table(rows=1, cols=3)
    borders(table)
    width(table, [0.7, 2.3, 3.5])
    for i, h in enumerate(("步", "发生位置", "意义")):
        cell = table.rows[0].cells[i]
        cell.text = h
        shade(cell, "E8EEF5")
        for r in cell.paragraphs[0].runs:
            r.bold = True
    for idx, (where, why) in enumerate(rows, 1):
        cells = table.add_row().cells
        cells[0].text = str(idx)
        cells[1].text = where
        cells[2].text = why
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for cell in cells:
            margins(cell)
    doc.add_paragraph()


def code(doc, text):
    table = doc.add_table(rows=1, cols=1)
    borders(table, "CBD5E1")
    width(table, [6.5])
    cell = table.cell(0, 0)
    shade(cell, "F8FAFC")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    run_font(r, 8.5, False, None, True)
    doc.add_paragraph()


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def build():
    doc = Document()
    style(doc)
    title(doc)

    doc.add_heading("1. 本次更新解决了什么", level=1)
    callout(
        doc,
        "一句话",
        "之前敌人的默认属性更像是直接从 CharacterBase 的三个 GE 来；现在变成了“按敌人种类查配置表，再应用默认属性”的数据驱动结构。",
    )
    tri(
        doc,
        ("改动点", "之前的思路", "现在的思路"),
        [
            ("Primary Attributes", "角色基类直接持有 DefaultPrimaryAttributes。", "每个 ECharacterClass 在 DA_CharacterClassInfo 里配置自己的 PrimaryAttributes GE。"),
            ("Secondary / Vital", "角色基类直接持有 DefaultSecondaryAttributes / DefaultVitalAttributes。", "放到 UCharacterClassInfo 的 Common Class Defaults 中，所有敌人类型共用。"),
            ("敌人差异", "蓝图或类里单独配 GE，扩展时容易散。", "敌人只配 CharacterClass 和 Level，由 Library 根据类型统一初始化。"),
            ("初始化入口", "AAuraCharacterBase::InitializeDefaultAttributes 直接 ApplyEffectToSelf。", "AAuraEnemy override InitializeDefaultAttributes，转调 UAuraAbilitySystemLibrary。"),
        ],
    )

    doc.add_heading("2. 新增结构职责地图", level=1)
    kv(
        doc,
        [
            ("ECharacterClass", "敌人/角色职业种类枚举，当前包含 Elementalist、Warrior、Ranger。它是 Map 的 Key，也是敌人蓝图里选择类型的字段。"),
            ("FCharacterClassDefaultInfo", "每个职业独有的一包默认信息，目前只放 PrimaryAttributes。后续可以扩展 XPReward、技能表、AI 行为等。"),
            ("UCharacterClassInfo", "DataAsset，总配置表。内部有 CharacterClassInformation Map，以及共用 SecondaryAttributes / VitalAttributes。"),
            ("DA_CharacterClassInfo", "蓝图数据资产实例，在编辑器里填具体 GE。GameMode 指向它，运行时 C++ 从它查配置。"),
            ("AAuraGameModeBase", "关卡规则层持有 CharacterClassInfo。Library 通过 UGameplayStatics::GetGameMode 取到它。"),
            ("UAuraAbilitySystemLibrary", "公共初始化函数 InitializeDefaultAttributes，把“查表 + Make Spec + Apply GE”封装起来。"),
            ("AAuraEnemy", "保存 CharacterClass 和 Level，初始化 ASC 后调用 Library 给自己套默认属性。"),
        ],
    )

    doc.add_heading("3. 数据定义：CharacterClassInfo", level=1)
    code(
        doc,
        """UENUM(BlueprintType)
enum class ECharacterClass : uint8
{
    Elementalist,
    Warrior,
    Ranger
};

USTRUCT(BlueprintType)
struct FCharacterClassDefaultInfo
{
    GENERATED_BODY()

    UPROPERTY(EditDefaultsOnly, Category = "Class Defaults")
    TSubclassOf<UGameplayEffect> PrimaryAttributes;
};""",
    )
    callout(
        doc,
        "怎么理解",
        "ECharacterClass 是“敌人属于哪一类”；FCharacterClassDefaultInfo 是“这一类敌人的专属默认配置”。现在它只装 PrimaryAttributes，所以 Warrior、Ranger、Elementalist 可以有不同力量/智力/韧性/活力基础值。",
    )
    code(
        doc,
        """UPROPERTY(EditDefaultsOnly, Category = "Character Class Defaults")
TMap<ECharacterClass, FCharacterClassDefaultInfo> CharacterClassInformation;

UPROPERTY(EditDefaultsOnly, Category = "Common Class Defaults")
TSubclassOf<UGameplayEffect> SecondaryAttributes;

UPROPERTY(EditDefaultsOnly, Category = "Common Class Defaults")
TSubclassOf<UGameplayEffect> VitalAttributes;""",
    )
    callout(
        doc,
        "为什么 Secondary / Vital 不放进每个职业里",
        "因为 SecondaryAttributes 通常由 PrimaryAttributes 计算，VitalAttributes 通常负责把 Health / Mana 设置到最大值。它们的计算规则共用即可，真正差异来自 PrimaryAttributes。",
    )

    doc.add_heading("4. 蓝图数据资产怎么填", level=1)
    tri(
        doc,
        ("字段", "编辑器里应该填什么", "意义"),
        [
            ("CharacterClassInformation", "为 Elementalist / Warrior / Ranger 每个 Key 添加一行。", "告诉系统每种敌人要用哪套 Primary GE。"),
            ("PrimaryAttributes", "例如 Warrior 对应 GE_EnemyWarriorPrimaryAttributes。", "决定这个种类的基础力量、智力、韧性、活力等。"),
            ("SecondaryAttributes", "例如 GE_EnemySecondaryAttributes。", "共用二级属性规则，如护甲、暴击、最大生命等。"),
            ("VitalAttributes", "例如 GE_EnemyVitalAttributes。", "最终把 Health / Mana 等生命资源初始化到合理值。"),
            ("GameMode.CharacterClassInfo", "在 BP_AuraGameMode 或当前 GameMode 默认值里指向 DA_CharacterClassInfo。", "否则 Library 从 GameMode 取不到配置表，敌人不会初始化默认属性。"),
            ("Enemy.CharacterClass", "在 BP_EnemyBase 或子类敌人里选择 Warrior / Ranger / Elementalist。", "同一个 C++ 敌人类可以靠这个字段变成不同默认属性类型。"),
        ],
    )

    doc.add_heading("5. 初始化运行流程", level=1)
    flow(
        doc,
        "敌人 BeginPlay 到属性生效",
        [
            ("AAuraEnemy::BeginPlay", "调用 InitAbilityActorInfo，开始初始化敌人的 ASC 和 Avatar/Owner。"),
            ("AAuraEnemy::InitAbilityActorInfo", "AbilitySystemComponent->InitAbilityActorInfo(this, this)，然后调用 AbilityActorInfoSet。"),
            ("AAuraEnemy::InitializeDefaultAttributes", "Enemy override 基类函数，不再使用 CharacterBase 的三个默认 GE 字段，而是调用 Library。"),
            ("UAuraAbilitySystemLibrary::InitializeDefaultAttributes", "通过 WorldContextObject 找 GameMode，再拿 GameMode->CharacterClassInfo。"),
            ("UCharacterClassInfo::GetClassDefaultInfo", "用敌人的 CharacterClass 作为 Key，从 CharacterClassInformation 里 FindChecked 对应配置。"),
            ("ASC->MakeEffectContext", "为每个默认 GE 创建 Context，并 AddSourceObject(AvatarActor)。"),
            ("ASC->MakeOutgoingSpec", "按 Level 创建 Primary / Secondary / Vital 三个 Spec。"),
            ("ASC->ApplyGameplayEffectSpecToSelf", "把三个默认 GE 应用到敌人自己的 ASC，AttributeSet 数值被写入或计算。"),
            ("敌人血条广播", "属性初始化后，BeginPlay 后半段广播当前 Health / MaxHealth，头顶血条得到初始值。"),
        ],
    )
    code(
        doc,
        """void AAuraEnemy::InitializeDefaultAttributes()
{
    UAuraAbilitySystemLibrary::InitializeDefaultAttributes(
        this,
        CharacterClass,
        Level,
        AbilitySystemComponent
    );
}""",
    )

    doc.add_heading("6. Library 里三段 GE 的顺序", level=1)
    tri(
        doc,
        ("顺序", "应用的 GE", "为什么这个顺序合理"),
        [
            ("1", "ClassDefaultInfo.PrimaryAttributes", "先写入职业差异的主属性。比如 Warrior 的 Vigor 更高，Elementalist 的 Intelligence 更高。"),
            ("2", "CharacterClassInfo->SecondaryAttributes", "二级属性往往基于主属性或 MMC 计算，所以要在主属性之后。"),
            ("3", "CharacterClassInfo->VitalAttributes", "生命/魔力当前值通常依赖 MaxHealth / MaxMana，最后初始化更稳。"),
        ],
        widths=(0.8, 2.4, 3.3),
    )
    callout(
        doc,
        "记忆方式",
        "先定体质和职业底子 Primary，再算派生战斗属性 Secondary，最后把当前血蓝 Vital 灌满。顺序反过来就容易出现当前生命先算了，最大生命后变了的别扭情况。",
    )

    doc.add_heading("7. 这套设计和玩家默认属性的区别", level=1)
    tri(
        doc,
        ("对象", "当前初始化方式", "适合原因"),
        [
            ("玩家角色", "仍可使用 AAuraCharacterBase 的 DefaultPrimary / Secondary / Vital 字段，或后续也迁移到 ClassInfo。", "玩家通常有职业选择、存档、升级等额外系统，后面可以单独设计。"),
            ("敌人角色", "AAuraEnemy override InitializeDefaultAttributes，使用 CharacterClassInfo。", "敌人种类多，用数据资产统一管理比每个蓝图手填三套 GE 更清楚。"),
            ("共用部分", "最终都还是 MakeOutgoingSpec + ApplyGameplayEffectSpecToSelf。", "GAS 入口没有变，只是“GE 从哪里来”变得数据驱动。"),
        ],
    )

    doc.add_heading("8. 常见错误与排查", level=1)
    bullets(
        doc,
        [
            "如果敌人出生后属性是 0 或血条异常，先查 GameMode 的 CharacterClassInfo 是否指向 DA_CharacterClassInfo。",
            "如果崩在 FindChecked，通常是 CharacterClassInformation 里没有对应的 Key，比如敌人选了 Ranger，但 DataAsset 没填 Ranger 行。",
            "如果编译报 UAbilitySystemComponent 未定义类型，说明 cpp 里调用 ASC 方法时缺少 #include \"AbilitySystemComponent.h\"。",
            "如果 UAuraAbilitySystemLibrary 不是类或命名空间，说明使用它的 cpp 没 include \"AbilitySystem/AuraAbilitySystemLibrary.h\"。",
            "如果 Secondary / Vital 没生效，检查 DA_CharacterClassInfo 的 Common Class Defaults 是否为空。",
            "如果 MaxHealth 算不对，回头查 SecondaryAttributes GE 里的 MMC 是否依赖 Context 的 SourceObject；现在 Library 已经 AddSourceObject(AvatarActor)。",
            "如果敌人蓝图改了 CharacterClass 但 PIE 不变化，确认实际生成的是你改过的那个蓝图子类，不是另一个敌人 BP。",
        ],
    )

    doc.add_heading("9. 后续可扩展方向", level=1)
    bullets(
        doc,
        [
            "把 FCharacterClassDefaultInfo 扩展为包含 StartupAbilities，让不同敌人种类自动获得不同技能。",
            "加入 XPReward、LootTable、AIBehavior、DamageTypeResistance 等字段，让敌人种类真正成为配置中心。",
            "把玩家职业也迁移到类似 ClassInfo 或 PlayerClassInfo，保持默认属性初始化入口统一。",
            "为 UCharacterClassInfo::GetClassDefaultInfo 加更友好的错误日志，替代 FindChecked 直接崩溃的学习期体验。",
            "把默认属性初始化包装成可复用的测试清单：生成敌人后打印 Class、Level、Health、MaxHealth、Primary 属性。",
        ],
    )

    doc.add_heading("10. 心智模型总结", level=1)
    callout(
        doc,
        "DataAsset 是菜单",
        "DA_CharacterClassInfo 就像一张菜单：Warrior 点哪套主属性，Ranger 点哪套主属性，都写在菜单里。",
    )
    callout(
        doc,
        "GameMode 是餐厅经理",
        "AAuraGameModeBase 持有这张菜单。Library 不自己猜配置，而是向 GameMode 要当前关卡认可的那张菜单。",
    )
    callout(
        doc,
        "Enemy 是点单的人",
        "AAuraEnemy 只说“我是 Warrior，等级 1”。Library 根据这句话去菜单上找对应 GE，再按顺序把属性应用到 ASC。这样扩展敌人时，C++ 不用频繁改。"
    )

    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
