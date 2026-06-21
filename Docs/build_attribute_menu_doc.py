from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "Docs/Aura_AttributeMenu_Technical_Guide.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="B8C7D9", size="4"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("Aura 属性菜单技术文档")
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Attribute Menu / WidgetController / GameplayTag 广播更新流程")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string("555555")

    meta = doc.add_table(rows=4, cols=2)
    set_table_borders(meta, color="D7DEE8", size="4")
    set_table_width(meta, [1.65, 4.85])
    rows = [
        ("项目", "F:\\ueprojiect\\Aura"),
        ("主题", "属性菜单 UI 从 C++ 数据到蓝图行控件的完整链路"),
        ("当前状态", "Primary Attributes 已完成，Secondary Attributes 已接入 TagsToAttributes，菜单可正常使用"),
        ("建议用途", "学习复盘、后续上下文压缩后的代码导航、排查属性菜单显示问题"),
    ]
    for row, (label, value) in zip(meta.rows, rows):
        set_cell_shading(row.cells[0], "E8EEF5")
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[1].paragraphs[0].add_run(value)


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_borders(table, color="B8C7D9", size="4")
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string("1F3A5F")
    p.add_run("  " + body)
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.add_run(item)


def add_code(doc, code):
    table = doc.add_table(rows=1, cols=1)
    set_table_borders(table, color="CBD5E1", size="4")
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F8FAFC")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(code)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9)


def add_kv_table(doc, rows, widths=(1.85, 4.65)):
    table = doc.add_table(rows=1, cols=2)
    set_table_borders(table, color="D7DEE8", size="4")
    set_table_width(table, list(widths))
    hdr = table.rows[0].cells
    hdr[0].text = "位置 / 名称"
    hdr[1].text = "作用"
    for cell in hdr:
        set_cell_shading(cell, "E8EEF5")
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
        for c in cells:
            set_cell_margins(c)
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()


def build_doc():
    doc = Document()
    style_doc(doc)
    add_title(doc)

    doc.add_heading("1. 核心结论", level=1)
    add_callout(
        doc,
        "一句话版本",
        "属性菜单不是每一行自己去读 AttributeSet，而是 MenuWidgetController 统一广播 FAuraAttributeInfo；每个 Row 收到广播后，用自己的 AttributeTag 和 Info.AttributeTag 做精确匹配，匹配成功才更新该行。",
    )
    add_bullets(
        doc,
        [
            "C++ 负责准备数据源：GameplayTag、FGameplayAttribute getter、AttributeInfo DataAsset、ASC 属性变化回调。",
            "MenuWidgetController 负责把 AttributeSet 中的属性转换成 UI 需要的 FAuraAttributeInfo 并广播。",
            "蓝图 Row 负责监听广播，并用本行配置的 AttributeTag 判断自己是不是该更新。",
            "父类 WBP_TextValueRow 放通用显示逻辑，子类 WBP_TextValueButtonRow 只扩展按钮相关逻辑，子类重写事件时要 Call Parent Function。",
        ],
    )

    doc.add_heading("2. 涉及文件总览", level=1)
    add_kv_table(
        doc,
        [
            ("Source/Aura/Public/AuraGameplayTags.h", "声明项目中要使用的 GameplayTag 成员变量，例如 Attributes_Primary_Strength、Attributes_Secondary_Armor。"),
            ("Source/Aura/Private/AuraGameplayTags.cpp", "在 InitializeNativeGameplayTags 中注册 Native Gameplay Tags，把字符串标签变成运行时可用的 FGameplayTag。"),
            ("Source/Aura/Public/AbilitySystem/AuraAttributeSet.h", "声明 GameplayAttributeData 属性、ATTRIBUTE_ACCESSORS、TagsToAttributes 映射表以及 OnRep 函数。"),
            ("Source/Aura/Private/AbilitySystem/AuraAttributeSet.cpp", "在构造函数里把 GameplayTag 映射到对应属性 getter；处理复制、PreAttributeChange、PostGameplayEffectExecute。"),
            ("Source/Aura/Public/AbilitySystem/Data/AttributeInfo.h", "定义 FAuraAttributeInfo，包含 AttributeTag、AttributeName、AttributeDescription、Attribute 数值。"),
            ("Source/Aura/Private/AbilitySystem/Data/AttributeInfo.cpp", "FindAttributeInfoForTag 根据 Tag 从 DataAsset 数组里找到 UI 展示信息。"),
            ("Source/Aura/Public/UI/WidgetController/MenuWidgetController.h", "声明 AttributeInfoDelegate、BroadcastInitialValues、BindCallbacksToDependencies、BroadcastAttributeInfo。"),
            ("Source/Aura/Private/UI/WidgetController/MenuWidgetController.cpp", "属性菜单广播的核心实现：初始广播和属性变化回调广播。"),
            ("Source/Aura/Private/UI/HUD/AuraHUD.cpp", "创建并缓存 AttributeMenuWidgetController，设置 WidgetControllerParams，绑定依赖。"),
            ("Source/Aura/Private/AbilitySystem/AuraAbilitySystemLibrary.cpp", "给蓝图提供 GetAttributeMenuWidgetController 入口，从 PC/HUD/PS 找到菜单 Controller。"),
            ("WBP_TextValueRow", "属性行父类，存放通用监听、Tag 匹配、Label/Value 更新逻辑。"),
            ("WBP_TextValueButtonRow", "属性行子类，继承父类通用更新逻辑，并额外放按钮逻辑。"),
        ],
    )

    doc.add_heading("3. C++ 数据结构层", level=1)
    doc.add_heading("3.1 GameplayTag 注册", level=2)
    doc.add_paragraph(
        "FAuraGameplayTags 是项目自己的 Tag 集合。h 文件里声明变量，cpp 文件里用 UGameplayTagsManager::Get().AddNativeGameplayTag 注册。"
    )
    add_code(
        doc,
        "GameplayTags.Attributes_Secondary_Armor = UGameplayTagsManager::Get().AddNativeGameplayTag(\n"
        "    FName(\"Attributes.Secondary.Armor\"),\n"
        "    FString(\"reduce damage,rise the chance to block\")\n"
        ");",
    )
    add_bullets(
        doc,
        [
            "h 文件中的 FGameplayTag Attributes_Secondary_Armor 只是一个变量槽位。",
            "cpp 中 AddNativeGameplayTag 才是真正把 Attributes.Secondary.Armor 这个字符串注册进 GameplayTags 系统。",
            "蓝图、DataAsset、Row Widget、C++ Map 必须使用同一套 Tag；只要有一处不一致，匹配就会失败。",
        ],
    )

    doc.add_heading("3.2 AttributeSet 与 TagsToAttributes", level=2)
    doc.add_paragraph(
        "UAuraAttributeSet 里每个属性用 FGameplayAttributeData 保存，ATTRIBUTE_ACCESSORS 会生成 GetHealthAttribute、GetHealth、SetHealth、InitHealth 等辅助函数。"
    )
    add_code(
        doc,
        "UPROPERTY(BlueprintReadOnly, ReplicatedUsing = OnRep_Armor, Category = \"Secondary Attribute\")\n"
        "FGameplayAttributeData Armor;\n"
        "ATTRIBUTE_ACCESSORS(UAuraAttributeSet, Armor);",
    )
    doc.add_paragraph(
        "TagsToAttributes 是属性菜单能自动遍历所有属性的关键。它把 UI 用的 GameplayTag 映射到 GAS 用的 FGameplayAttribute getter。"
    )
    add_code(
        doc,
        "TMap<FGameplayTag, TStaticFuncPtr<FGameplayAttribute()>> TagsToAttributes;\n\n"
        "TagsToAttributes.Add(GameplayTags.Attributes_Primary_Strength, GetStrengthAttribute);\n"
        "TagsToAttributes.Add(GameplayTags.Attributes_Secondary_Armor, GetArmorAttribute);\n"
        "TagsToAttributes.Add(GameplayTags.Attributes_Secondary_MaxHealth, GetMaxHealthAttribute);",
    )
    add_callout(
        doc,
        "重要",
        "属性菜单广播的源头就是 TagsToAttributes。某个属性没有 Add 到这个 Map 里，MenuWidgetController 就不会遍历到它，也就不会广播给蓝图行。",
    )

    doc.add_heading("3.3 AttributeInfo DataAsset", level=2)
    doc.add_paragraph(
        "FAuraAttributeInfo 是 UI 需要的展示数据结构，不是 GAS 原生属性。它把属性 Tag、显示名字、描述文本和当前数值打包到一起。"
    )
    add_code(
        doc,
        "USTRUCT(BlueprintType)\n"
        "struct FAuraAttributeInfo\n"
        "{\n"
        "    GENERATED_BODY()\n"
        "    FGameplayTag AttributeTag;\n"
        "    FText AttributeName;\n"
        "    FText AttributeDescription;\n"
        "    float Attribute = 0.f;\n"
        "};",
    )
    doc.add_paragraph(
        "UAttributeInfo::FindAttributeInfoForTag 会遍历 DataAsset 中的 AttributeInformation 数组，找到 AttributeTag 精确匹配的那一项。"
    )
    add_bullets(
        doc,
        [
            "DataAsset 里必须有 Primary 和 Secondary 的所有条目。",
            "每一项的 AttributeTag 要和 FAuraGameplayTags 注册出来的 Tag 完全一致。",
            "如果返回 None，蓝图 Row 的 Matches Tag Exact 会失败。",
            "调试时可以临时传 true：FindAttributeInfoForTag(Pair.Key, true)，缺项会打日志。",
        ],
    )

    doc.add_heading("4. WidgetController 层", level=1)
    doc.add_heading("4.1 AuraHUD 创建菜单 Controller", level=2)
    doc.add_paragraph(
        "AAuraHUD::GetAttributeMenuWidgetController 负责创建并缓存 UMenuWidgetController。它不是每次打开菜单都重新 new 一个，而是第一次为空时创建，之后复用。"
    )
    add_code(
        doc,
        "if (AttributeMenuWidgetController == nullptr)\n"
        "{\n"
        "    AttributeMenuWidgetController = NewObject<UMenuWidgetController>(this, AttributeMenuWidgetControllerClass);\n"
        "    AttributeMenuWidgetController->SetWidgetControllerParams(WCParams);\n"
        "    AttributeMenuWidgetController->BindCallbacksToDependencies();\n"
        "}\n"
        "return AttributeMenuWidgetController;",
    )
    add_bullets(
        doc,
        [
            "SetWidgetControllerParams 把 PC、PS、ASC、AS 塞进 Controller。",
            "BindCallbacksToDependencies 建立 ASC 属性变化回调。",
            "BroadcastInitialValues 通常由蓝图菜单打开后主动调用，用于先显示当前属性值。",
        ],
    )

    doc.add_heading("4.2 AuraAbilitySystemLibrary 给蓝图找 Controller", level=2)
    doc.add_paragraph(
        "UAuraAbilitySystemLibrary::GetAttributeMenuWidgetController 是蓝图工具函数。蓝图不需要自己知道 PlayerState、ASC、AttributeSet 怎么找，只调用这个函数拿到菜单 Controller。"
    )
    add_numbered(
        doc,
        [
            "通过 UGameplayStatics::GetPlayerController(WorldContextObject, 0) 找到本地 PlayerController。",
            "从 PlayerController 的 HUD Cast 到 AAuraHUD。",
            "从 PlayerState 拿到 ASC 和 AttributeSet。",
            "组装 FWidgetControllerParams。",
            "调用 AuraHUD->GetAttributeMenuWidgetController 返回最终 Controller。",
        ],
    )

    doc.add_heading("4.3 MenuWidgetController 的初始广播", level=2)
    doc.add_paragraph(
        "UMenuWidgetController::BroadcastInitialValues 用于菜单刚打开时，把所有当前属性值广播一次，让 UI 行先显示初始值。"
    )
    add_code(
        doc,
        "void UMenuWidgetController::BroadcastInitialValues()\n"
        "{\n"
        "    UAuraAttributeSet* AS = CastChecked<UAuraAttributeSet>(AttributeSet);\n"
        "    check(AttributeInfo);\n\n"
        "    for (auto& Pair : AS->TagsToAttributes)\n"
        "    {\n"
        "        BroadcastAttributeInfo(Pair.Key, Pair.Value());\n"
        "    }\n"
        "}",
    )

    doc.add_heading("4.4 MenuWidgetController 的变化回调", level=2)
    doc.add_paragraph(
        "UMenuWidgetController::BindCallbacksToDependencies 会给 TagsToAttributes 里的每个属性都注册一个 ASC 变化监听。只要属性变化，Lambda 就触发，再广播给 UI。"
    )
    add_code(
        doc,
        "for (auto& Pair : AS->TagsToAttributes)\n"
        "{\n"
        "    AbilitySystemComponent->GetGameplayAttributeValueChangeDelegate(Pair.Value()).AddLambda(\n"
        "        [this, Pair](const FOnAttributeChangeData& Data)\n"
        "        {\n"
        "            BroadcastAttributeInfo(Pair.Key, Pair.Value());\n"
        "        }\n"
        "    );\n"
        "}",
    )
    add_callout(
        doc,
        "推荐小优化",
        "在变化回调里也可以让 BroadcastAttributeInfo 接收 Data.NewValue，避免再次从 AttributeSet 读取数值。不过你当前从 Attribute.GetNumericValue(AttributeSet) 读取，只要 AttributeSet 是正确对象，也能工作。",
    )

    doc.add_heading("4.5 BroadcastAttributeInfo 的职责", level=2)
    doc.add_paragraph(
        "BroadcastAttributeInfo 是你抽出来的复用函数，避免初始广播和变化回调里重复写 FindAttributeInfoForTag、取数值、Broadcast 这三步。"
    )
    add_code(
        doc,
        "void UMenuWidgetController::BroadcastAttributeInfo(const FGameplayTag& AttributeTag, const FGameplayAttribute& Attribute)\n"
        "{\n"
        "    FAuraAttributeInfo Info = AttributeInfo->FindAttributeInfoForTag(AttributeTag);\n"
        "    Info.Attribute = Attribute.GetNumericValue(AttributeSet);\n"
        "    AttributeInfoDelegate.Broadcast(Info);\n"
        "}",
    )
    add_bullets(
        doc,
        [
            "AttributeTag 用来从 DataAsset 找名字和描述。",
            "FGameplayAttribute 用来从 AttributeSet 取当前数值。",
            "AttributeInfoDelegate.Broadcast(Info) 会通知所有已经绑定这个委托的蓝图 Row。",
        ],
    )

    doc.add_heading("5. 蓝图 UI 层", level=1)
    doc.add_heading("5.1 WBP_TextValueRow 父类逻辑", level=2)
    doc.add_paragraph(
        "WBP_TextValueRow 应该放所有属性行都通用的逻辑：拿 MenuWidgetController、绑定 AttributeInfoDelegate、收到 Info 后匹配 Tag、更新 Label 和 Value。"
    )
    add_code(
        doc,
        "Event WidgetControllerSet\n"
        "    -> Get Attribute Menu Widget Controller\n"
        "    -> Bind Event to Attribute Info Delegate\n\n"
        "AttributeInfoDelegate_Event(Info)\n"
        "    -> Break Aura Attribute Info\n"
        "    -> Matches Tag Exact(Row.AttributeTag, Info.AttributeTag)\n"
        "    -> Branch\n"
        "        True -> Set Label Text(Info.AttributeName)\n"
        "             -> Set Numerical Value(Info.Attribute)",
    )
    add_bullets(
        doc,
        [
            "Event Construct 是控件出生时触发，此时 WidgetController 可能还没设置好。",
            "Event WidgetControllerSet 是 C++/父类 SetWidgetController 后触发，更适合绑定 Controller 相关回调。",
            "每个 Row 自己有一个 AttributeTag 变量，表示这一行想显示哪个属性。",
            "广播不是只发给某一行，而是所有绑定者都收到；每一行自己判断是不是该更新。",
        ],
    )

    doc.add_heading("5.2 WBP_TextValueButtonRow 子类逻辑", level=2)
    doc.add_paragraph(
        "WBP_TextValueButtonRow 是 WBP_TextValueRow 的子类。它可以复用父类的标签匹配和数值更新逻辑，只额外添加按钮、点击、可用状态等逻辑。"
    )
    add_callout(
        doc,
        "关键点",
        "如果子类也重写 Event WidgetControllerSet，必须 Add Call to Parent Function。否则父类里的绑定广播逻辑会被覆盖，导致子类行收不到属性更新。",
    )
    add_code(
        doc,
        "WBP_TextValueButtonRow.Event WidgetControllerSet\n"
        "    -> Parent: WidgetControllerSet\n"
        "    -> 子类自己的按钮逻辑",
    )

    doc.add_heading("6. 完整运行流程", level=1)
    add_numbered(
        doc,
        [
            "角色初始化 ASC 和 AttributeSet，默认 GE 初始化 Primary、Secondary、Vital 属性。",
            "玩家打开属性菜单蓝图。",
            "蓝图调用 UAuraAbilitySystemLibrary::GetAttributeMenuWidgetController。",
            "函数通过 PC -> HUD -> PlayerState 找到 ASC 和 AttributeSet，然后让 AuraHUD 返回菜单 Controller。",
            "AuraHUD 第一次创建 UMenuWidgetController，并调用 SetWidgetControllerParams 和 BindCallbacksToDependencies。",
            "各个属性 Row 的 Event WidgetControllerSet 绑定 AttributeInfoDelegate。",
            "菜单蓝图调用 BroadcastInitialValues。",
            "MenuWidgetController 遍历 TagsToAttributes，把每个属性打包成 FAuraAttributeInfo 并广播。",
            "所有 Row 都收到广播，但只有 Row.AttributeTag 与 Info.AttributeTag 匹配的那一行更新。",
            "后续属性变化时，ASC 的 GetGameplayAttributeValueChangeDelegate 触发 Lambda，再次广播对应属性 Info，UI 自动刷新。",
        ],
    )

    doc.add_heading("7. Primary 与 Secondary 接入清单", level=1)
    table = doc.add_table(rows=1, cols=4)
    set_table_borders(table, color="D7DEE8", size="4")
    set_table_width(table, [1.6, 1.8, 1.6, 1.5])
    hdr = table.rows[0].cells
    for idx, text in enumerate(["步骤", "Primary 示例", "Secondary 示例", "检查点"]):
        hdr[idx].text = text
        set_cell_shading(hdr[idx], "E8EEF5")
        for r in hdr[idx].paragraphs[0].runs:
            r.bold = True
    rows = [
        ("声明 Tag 变量", "Attributes_Primary_Strength", "Attributes_Secondary_Armor", "AuraGameplayTags.h"),
        ("注册 Tag 字符串", "Attributes.Primary.Strength", "Attributes.Secondary.Armor", "AuraGameplayTags.cpp"),
        ("声明 AttributeData", "Strength", "Armor", "AuraAttributeSet.h"),
        ("添加 Accessors", "GetStrengthAttribute", "GetArmorAttribute", "ATTRIBUTE_ACCESSORS"),
        ("加入 TagsToAttributes", "Tag -> GetStrengthAttribute", "Tag -> GetArmorAttribute", "AuraAttributeSet.cpp 构造函数"),
        ("DataAsset 添加信息", "力量 / 描述", "护甲 / 描述", "AttributeInfo DataAsset"),
        ("Row 配置 Tag", "行 AttributeTag = Strength", "行 AttributeTag = Armor", "WBP_AttributeMenu"),
        ("默认 GE 设置值", "DefaultPrimaryAttributes", "DefaultSecondaryAttributes", "角色蓝图和 GE 蓝图"),
    ]
    for row_data in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row_data):
            cells[idx].text = text
            set_cell_margins(cells[idx])
    doc.add_paragraph()

    doc.add_heading("8. 常见问题与排查", level=1)
    add_kv_table(
        doc,
        [
            ("UI 行不更新", "先打印 Row.AttributeTag 和 Info.AttributeTag。两者必须 Matches Tag Exact 成功。"),
            ("打印出 None", "通常是 Row 的 AttributeTag 没设置，或者 AttributeInfo DataAsset 没有对应条目。"),
            ("Primary 可以，Secondary 不行", "确认 Secondary 已加入 TagsToAttributes，DataAsset 有 Secondary 项，DefaultSecondaryAttributes GE 有实际赋值。"),
            ("父类 Row 写了逻辑，子类不执行", "子类重写 Event WidgetControllerSet 后没有 Call Parent Function。"),
            ("编译报 FGameplayTag class/struct 冲突", "不要乱写 class FGameplayTag; 前向声明。包含 GameplayTagContainer.h。FGameplayAttribute 需要 AttributeSet.h。"),
            ("BroadcastAttributeInfo 不接受 2 个参数", "检查 MenuWidgetController.h 声明和 MenuWidgetController.cpp 定义是否完全一致，并确保类型头文件已包含。"),
            ("Live Coding 导致 VS 编译失败", "关闭 UE 编辑器后从 VS 编译，或者在 UE 中按 Ctrl+Alt+F11。改 UPROPERTY/UFUNCTION/头文件时建议关编辑器完整编译。"),
            ("变化后 UI 不刷新", "确认 BindCallbacksToDependencies 已执行，ASC 是正确对象，属性变化确实走了 GetGameplayAttributeValueChangeDelegate。"),
        ],
    )

    doc.add_heading("9. 推荐调试日志", level=1)
    doc.add_paragraph("当属性菜单又开始迷糊时，可以临时加这些日志，一步步确认数据是否经过每一层。")
    add_code(
        doc,
        "UE_LOG(LogTemp, Warning, TEXT(\"MapNum: %d\"), AS->TagsToAttributes.Num());\n"
        "UE_LOG(LogTemp, Warning, TEXT(\"Pair.Key: %s\"), *Pair.Key.ToString());\n\n"
        "FAuraAttributeInfo Info = AttributeInfo->FindAttributeInfoForTag(Pair.Key, true);\n"
        "UE_LOG(LogTemp, Warning, TEXT(\"Info.AttributeTag: %s\"), *Info.AttributeTag.ToString());\n"
        "UE_LOG(LogTemp, Warning, TEXT(\"Info.Attribute: %f\"), Info.Attribute);",
    )
    add_bullets(
        doc,
        [
            "MapNum 不对：查 UAuraAttributeSet 构造函数的 TagsToAttributes.Add。",
            "Pair.Key 对，但 Info.AttributeTag 是 None：查 AttributeInfo DataAsset。",
            "Info 正确但 Row 不更新：查蓝图 Row 的 AttributeTag 和 Matches Tag Exact。",
            "初始显示但变化不刷新：查 BindCallbacksToDependencies 和 ASC 属性变化委托。",
        ],
    )

    doc.add_heading("10. 维护建议", level=1)
    add_bullets(
        doc,
        [
            "新增属性时按“Tag 声明 -> Tag 注册 -> AttributeData -> Accessors -> TagsToAttributes -> DataAsset -> Row Tag -> GE 值”这个顺序做。",
            "把通用显示逻辑放在 WBP_TextValueRow，避免每个子类复制一份绑定和匹配逻辑。",
            "C++ 的 BroadcastAttributeInfo 保持小而纯：找 Info、填数值、广播。",
            "蓝图里如果出现重复绑定导致多次打印，可以检查 Event WidgetControllerSet 是否被多次调用，必要时加 DoOnce 或解绑策略。",
            "Secondary 属性如果由 MMC 计算，确认 DefaultSecondaryAttributes 在 Primary 之后应用，否则计算依赖的 Primary 可能还没初始化。",
        ],
    )

    doc.add_heading("11. 心智模型", level=1)
    add_code(
        doc,
        "GameplayTag: UI 和属性之间的名字桥梁\n"
        "FGameplayAttribute: GAS 里真正能取数值的属性描述\n"
        "TagsToAttributes: Tag -> Attribute getter 的映射表\n"
        "AttributeInfo DataAsset: Tag -> UI 文案的配置表\n"
        "MenuWidgetController: 把 GAS 数据翻译成 UI 能用的信息并广播\n"
        "Row Widget: 每一行自己用 AttributeTag 过滤广播，只更新自己",
    )
    add_callout(
        doc,
        "最终记法",
        "C++ 广播很多属性；每个 Row 都能听见；但 Row 只认自己的 AttributeTag。这个模式让属性菜单可以扩展很多行，而不需要为每个属性单独写一套 UI 绑定。",
    )

    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Aura Attribute Menu Technical Guide")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string("666666")

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
