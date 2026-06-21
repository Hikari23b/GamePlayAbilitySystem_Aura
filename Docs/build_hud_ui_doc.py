from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT = "Docs/Aura_HUD_WidgetController_UI_Guide.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(10)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, header, True)
        set_cell_shading(cell, "E8EEF5")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if widths:
            cell.width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths:
                cells[i].width = Inches(widths[i])
    doc.add_paragraph()
    return table


def add_code(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(30, 30, 30)


def add_note(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.bold = True
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(11)
styles["Normal"].paragraph_format.space_after = Pt(6)
styles["Normal"].paragraph_format.line_spacing = 1.25

for style_name, size, color in [
    ("Heading 1", 16, "2E74B5"),
    ("Heading 2", 13, "2E74B5"),
    ("Heading 3", 12, "1F4D78"),
]:
    style = styles[style_name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(12)
    style.paragraph_format.space_after = Pt(6)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Aura HUD / WidgetController UI 技术文档")
run.bold = True
run.font.name = "Calibri"
run.font.size = Pt(20)
run.font.color.rgb = RGBColor.from_string("0B2545")

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.add_run("范围：HUD、Widget、WidgetController、GAS 属性回调；暂不包含 GameplayEffect 细节。")
subtitle.paragraph_format.space_after = Pt(12)

add_note(doc, "核心记忆：Character 负责启动；HUD 负责装配；WidgetController 负责数据；Widget 蓝图负责显示。")

doc.add_heading("1. 整体结构", level=1)
doc.add_paragraph("这一套 UI 结构不是让 Widget 自己到处找数据，而是由 HUD 创建 UI 和 Controller，再由 WidgetController 把 AbilitySystemComponent 与 AttributeSet 中的数据广播给蓝图。")

add_code(
    doc,
    "AAuraCharacter::InitAbilityActorInfo()\n"
    "    -> AAuraHUD::InitOverlay(...)\n"
    "        -> CreateWidget(OverlayWidgetClass)\n"
    "        -> GetOverlayWidgetController(...)\n"
    "            -> NewObject<UOverlayWidgetController>(...)\n"
    "            -> SetWidgetControllerParams(...)\n"
    "            -> BindCallbacksToDependencies()\n"
    "        -> OverlayWidget->SetWidgetController(WidgetController)\n"
    "            -> UAuraUserWidget::WidgetControllerSet() [Blueprint Event]\n"
    "        -> WidgetController->BroadcastInitialValues()\n"
    "        -> AddToViewport()",
)

doc.add_heading("2. 文件职责总表", level=1)
add_table(
    doc,
    ["文件", "核心类/函数", "职责"],
    [
        ["Source/Aura/Private/Character/AuraCharacter.cpp", "AAuraCharacter::InitAbilityActorInfo", "GAS 初始化完成后，取得 PlayerController、PlayerState、ASC、AttributeSet，并调用 HUD 初始化 Overlay。"],
        ["Source/Aura/Public/UI/HUD/AuraHUD.h", "AAuraHUD 成员变量", "保存 OverlayWidgetClass、OverlayWidgetControllerClass、OverlayWidget、OverlayWidgetController。"],
        ["Source/Aura/Private/UI/HUD/AuraHUD.cpp", "InitOverlay", "创建 Overlay Widget，创建或获取 WidgetController，把 Controller 交给 Widget，并广播初始值。"],
        ["Source/Aura/Private/UI/HUD/AuraHUD.cpp", "GetOverlayWidgetController", "首次创建 OverlayWidgetController，设置参数，并绑定 GAS 属性变化回调。"],
        ["Source/Aura/Public/UI/WidgetController/AuraWidgetController.h", "FWidgetControllerParams", "把 PC、PS、ASC、AttributeSet 打包成一个参数结构体。"],
        ["Source/Aura/Private/UI/WidgetController/AuraWidgetController.cpp", "SetWidgetControllerParams", "把参数包中的对象保存到 WidgetController 成员变量。"],
        ["Source/Aura/Public/UI/WidgetController/OverlayWidgetController.h", "动态多播委托", "声明 Health、MaxHealth、Mana、MaxMana 的蓝图可绑定事件。"],
        ["Source/Aura/Private/UI/WidgetController/OverlayWidgetController.cpp", "BroadcastInitialValues", "UI 创建时广播当前属性值，让 UI 一开始就显示正确。"],
        ["Source/Aura/Private/UI/WidgetController/OverlayWidgetController.cpp", "BindCallbacksToDependencies", "绑定 ASC 的属性变化委托，Health/Mana 改变时通知 UI。"],
        ["Source/Aura/Private/UI/Widget/AuraUserWidget.cpp", "SetWidgetController", "Widget 保存 Controller，并触发蓝图事件 WidgetControllerSet。"],
        ["WBP_Overlay", "Event WidgetControllerSet", "把 Controller 继续传给 WBP_HealthGlobe、WBP_ManaGlobe 等子控件。"],
        ["WBP_HealthGlobe / WBP_ManaGlobe", "Event WidgetControllerSet", "Cast Controller，保存为 BPOverlayWidgetController，绑定对应属性变化事件，更新进度条。"],
    ],
    [2.1, 1.7, 2.7],
)

doc.add_heading("3. 启动入口：AAuraCharacter", level=1)
doc.add_paragraph("文件：Source/Aura/Private/Character/AuraCharacter.cpp")
doc.add_paragraph("函数：AAuraCharacter::InitAbilityActorInfo()")
doc.add_paragraph("这个函数是 UI 初始化链路的起点。PossessedBy 和 OnRep_PlayerState 都会调用它。")
add_code(
    doc,
    "AAuraPlayerState* AuraPlayerState = GetPlayerState<AAuraPlayerState>();\n"
    "AuraPlayerState->GetAbilitySystemComponent()->InitAbilityActorInfo(AuraPlayerState, this);\n"
    "AbilitySystemComponent = AuraPlayerState->GetAbilitySystemComponent();\n"
    "AttributeSet = AuraPlayerState->GetAttributeSet();\n\n"
    "if (AAuraPlayerController* AuraPlayerController = Cast<AAuraPlayerController>(GetController()))\n"
    "{\n"
    "    if (AAuraHUD* AuraHUD = Cast<AAuraHUD>(AuraPlayerController->GetHUD()))\n"
    "    {\n"
    "        AuraHUD->InitOverlay(AuraPlayerController, AuraPlayerState, AbilitySystemComponent, AttributeSet);\n"
    "    }\n"
    "}",
)
add_bullets(
    doc,
    [
        "这时 ASC 和 AttributeSet 已经可用，所以 HUD 可以安全创建 UI。",
        "Cast<AAuraPlayerController> 是为了拿到玩家专属 Controller，再通过 GetHUD 拿 AAuraHUD。",
        "如果 AuraHUD 为 nullptr，通常是 GameMode 的 HUD Class 没设成 BP_AuraHUD 或 AAuraHUD。",
    ],
)

doc.add_heading("4. HUD：装配工", level=1)
doc.add_paragraph("文件：Source/Aura/Private/UI/HUD/AuraHUD.cpp")

doc.add_heading("4.1 InitOverlay", level=2)
doc.add_paragraph("InitOverlay 负责把 Overlay Widget 和 OverlayWidgetController 装配起来。")
add_code(
    doc,
    "UUserWidget* Widget = CreateWidget<UUserWidget>(GetWorld(), OverlayWidgetClass);\n"
    "OverlayWidget = Cast<UAuraUserWidget>(Widget);\n\n"
    "const FWidgetControllerParams WidgetControllerParams(PC, PS, ASC, AS);\n"
    "UOverlayWidgetController* WidgetController = GetOverlayWidgetController(WidgetControllerParams);\n\n"
    "OverlayWidget->SetWidgetController(WidgetController);\n"
    "WidgetController->BroadcastInitialValues();\n"
    "Widget->AddToViewport();",
)
add_bullets(
    doc,
    [
        "CreateWidget 创建实际显示的 UI。",
        "OverlayWidget = Cast<UAuraUserWidget>(Widget) 是为了调用 UAuraUserWidget::SetWidgetController。",
        "SetWidgetController 会触发蓝图 Event WidgetControllerSet，使蓝图有机会绑定事件。",
        "BroadcastInitialValues 在蓝图绑定之后调用，这样初始 Health/Mana 能被 UI 接收到。",
        "AddToViewport 最后执行，把 UI 加到屏幕上。",
    ],
)

doc.add_heading("4.2 GetOverlayWidgetController", level=2)
doc.add_paragraph("这个函数负责创建并缓存 Controller。")
add_code(
    doc,
    "if (OverlayWidgetController == nullptr)\n"
    "{\n"
    "    OverlayWidgetController = NewObject<UOverlayWidgetController>(this, OverlayWidgetControllerClass);\n"
    "    OverlayWidgetController->SetWidgetControllerParams(WCParams);\n"
    "    OverlayWidgetController->BindCallbacksToDependencies();\n"
    "    return OverlayWidgetController;\n"
    "}\n"
    "return OverlayWidgetController;",
)
add_bullets(
    doc,
    [
        "第一次调用时创建 Controller，之后复用同一个 Controller。",
        "SetWidgetControllerParams 必须在 BindCallbacksToDependencies 之前，因为绑定回调需要 ASC 和 AttributeSet。",
        "OverlayWidgetControllerClass 必须在 BP_AuraHUD 里设置，通常设成 BP_OverlayWidgetController。",
    ],
)

doc.add_heading("5. WidgetController 父类：UAuraWidgetController", level=1)
doc.add_paragraph("文件：Source/Aura/Public/UI/WidgetController/AuraWidgetController.h")
doc.add_paragraph("文件：Source/Aura/Private/UI/WidgetController/AuraWidgetController.cpp")
add_table(
    doc,
    ["成员/函数", "作用"],
    [
        ["FWidgetControllerParams", "参数包，保存 PlayerController、PlayerState、AbilitySystemComponent、AttributeSet。"],
        ["SetWidgetControllerParams", "把参数包中的对象保存到 Controller 成员变量。"],
        ["BroadcastInitialValues", "虚函数；父类为空，子类重写。"],
        ["BindCallbacksToDependencies", "虚函数；父类为空，子类重写。"],
        ["PlayerController / PlayerState", "UI 需要的玩家上下文。"],
        ["AbilitySystemComponent", "GAS 事件源，用来监听属性变化。"],
        ["AttributeSet", "属性源，最终 Cast 成 UAuraAttributeSet 读取 Health/Mana。"],
    ],
    [2.0, 4.3],
)

doc.add_heading("6. OverlayWidgetController：属性广播核心", level=1)
doc.add_paragraph("文件：Source/Aura/Public/UI/WidgetController/OverlayWidgetController.h")
doc.add_paragraph("文件：Source/Aura/Private/UI/WidgetController/OverlayWidgetController.cpp")

doc.add_heading("6.1 委托声明", level=2)
doc.add_paragraph("OverlayWidgetController 声明了四个动态多播委托，蓝图可以绑定它们。")
add_code(
    doc,
    "DECLARE_DYNAMIC_MULTICAST_DELEGATE_OneParam(FOnHealthChangedSignature, float, NewHealth);\n"
    "DECLARE_DYNAMIC_MULTICAST_DELEGATE_OneParam(FOnMaxHealthChangedSignature, float, NewMaxHealth);\n"
    "DECLARE_DYNAMIC_MULTICAST_DELEGATE_OneParam(FOnManaChangedSignature, float, NewMana);\n"
    "DECLARE_DYNAMIC_MULTICAST_DELEGATE_OneParam(FOnMaxManaChangedSignature, float, NewMaxMana);",
)
add_code(
    doc,
    "UPROPERTY(BlueprintAssignable, Category = \"GAS|Attribute\")\n"
    "FOnHealthChangedSignature OnHealthChanged;\n\n"
    "UPROPERTY(BlueprintAssignable, Category = \"GAS|Attribute\")\n"
    "FOnMaxHealthChangedSignature OnMaxHealthChanged;\n\n"
    "UPROPERTY(BlueprintAssignable, Category = \"GAS|Attribute\")\n"
    "FOnManaChangedSignature OnManaChanged;\n\n"
    "UPROPERTY(BlueprintAssignable, Category = \"GAS|Attribute\")\n"
    "FOnMaxManaChangedSignature OnMaxManaChanged;",
)
add_note(doc, "注意：OnManaChanged 必须使用 FOnManaChangedSignature；OnMaxManaChanged 必须使用 FOnMaxManaChangedSignature。否则蓝图事件参数会继续显示 NewHealth。")

doc.add_heading("6.2 BroadcastInitialValues", level=2)
doc.add_paragraph("UI 刚创建时，属性已经有值，但如果不广播，蓝图可能还不知道初始 Health/Mana。")
add_code(
    doc,
    "const UAuraAttributeSet* AuraAttributeSet = CastChecked<UAuraAttributeSet>(AttributeSet);\n\n"
    "OnHealthChanged.Broadcast(AuraAttributeSet->GetHealth());\n"
    "OnMaxHealthChanged.Broadcast(AuraAttributeSet->GetMaxHealth());\n"
    "OnManaChanged.Broadcast(AuraAttributeSet->GetMana());\n"
    "OnMaxManaChanged.Broadcast(AuraAttributeSet->GetMaxMana());",
)
add_bullets(
    doc,
    [
        "CastChecked 是因为父类只保存 UAttributeSet，但需要 UAuraAttributeSet 的 GetHealth/GetMana 等访问器。",
        "如果 UI 百分比依赖 Max 值，可以先广播 MaxHealth/MaxMana，再广播 Health/Mana。",
        "BroadcastInitialValues 应在蓝图绑定事件之后调用。",
    ],
)

doc.add_heading("6.3 BindCallbacksToDependencies", level=2)
doc.add_paragraph("这个函数把 WidgetController 的回调绑定到 ASC 的属性变化委托上。")
add_code(
    doc,
    "AbilitySystemComponent->GetGameplayAttributeValueChangeDelegate(\n"
    "    AuraAttributeSet->GetHealthAttribute())\n"
    "    .AddUObject(this, &UOverlayWidgetController::HealthChanged);\n\n"
    "AbilitySystemComponent->GetGameplayAttributeValueChangeDelegate(\n"
    "    AuraAttributeSet->GetManaAttribute())\n"
    "    .AddUObject(this, &UOverlayWidgetController::ManaChanged);",
)
add_bullets(
    doc,
    [
        "Dependency 指的是 ASC 和 AttributeSet：Controller 依赖它们提供属性变化事件。",
        "Health 改变时调用 HealthChanged；Mana 改变时调用 ManaChanged。",
        "回调函数里把 Data.NewValue 再广播给蓝图。",
    ],
)

doc.add_heading("7. AuraUserWidget：Widget 接收 Controller", level=1)
doc.add_paragraph("文件：Source/Aura/Public/UI/Widget/AuraUserWidget.h")
doc.add_paragraph("文件：Source/Aura/Private/UI/Widget/AuraUserWidget.cpp")
add_code(
    doc,
    "void UAuraUserWidget::SetWidgetController(UObject* InWidgetController)\n"
    "{\n"
    "    WidgetController = InWidgetController;\n"
    "    WidgetControllerSet();\n"
    "}",
)
add_bullets(
    doc,
    [
        "WidgetController 是 UObject 指针，所以主 UI 和子 UI 都可以接收不同类型的 Controller。",
        "WidgetControllerSet 是 BlueprintImplementableEvent，C++ 不写实现。",
        "蓝图中的 Event WidgetControllerSet 就是在这里被触发。",
    ],
)

doc.add_heading("8. 蓝图侧：Overlay 与子控件", level=1)

doc.add_heading("8.1 WBP_Overlay", level=2)
doc.add_paragraph("WBP_Overlay 是主 Overlay。它从 C++ 获得 WidgetController 后，要把它继续传给子控件。")
add_code(
    doc,
    "Event WidgetControllerSet\n"
    "    -> WBP_HealthGlobe.SetWidgetController(WidgetController)\n"
    "    -> WBP_ManaGlobe.SetWidgetController(WidgetController)",
)

doc.add_heading("8.2 WBP_HealthGlobe", level=2)
doc.add_paragraph("HealthGlobe 收到 Controller 后，必须先 Cast，再 Set，再 Bind。")
add_code(
    doc,
    "Event WidgetControllerSet\n"
    "    -> Cast To BP_OverlayWidgetController (Object = WidgetController)\n"
    "    -> Set BPOverlayWidgetController\n"
    "    -> Bind Event to OnHealthChanged\n"
    "    -> Bind Event to OnMaxHealthChanged",
)

doc.add_heading("8.3 WBP_ManaGlobe", level=2)
doc.add_paragraph("ManaGlobe 同理，但绑定 Mana 事件。")
add_code(
    doc,
    "Event WidgetControllerSet\n"
    "    -> Cast To BP_OverlayWidgetController (Object = WidgetController)\n"
    "    -> Set BPOverlayWidgetController\n"
    "    -> Bind Event to OnManaChanged\n"
    "    -> Bind Event to OnMaxManaChanged",
)
add_note(doc, "不要用 Sequence 把 Cast/Set/Bind 分成互不相连的多条执行线。最稳的是 Cast -> Set -> Bind1 -> Bind2 串成一条线。")

doc.add_heading("9. 配置检查", level=1)
add_table(
    doc,
    ["位置", "必须配置"],
    [
        ["BP_AuraGameMode", "HUD Class = BP_AuraHUD；Player Controller Class = BP_AuraPlayerController；Player State Class = BP_AuraPlayerState。"],
        ["BP_AuraHUD", "OverlayWidgetClass = WBP_Overlay；OverlayWidgetControllerClass = BP_OverlayWidgetController。"],
        ["WBP_Overlay", "拥有 WBP_HealthGlobe 和 WBP_ManaGlobe 子控件，并在 Event WidgetControllerSet 里传 Controller。"],
        ["WBP_HealthGlobe", "Event WidgetControllerSet 中 Cast/Set 后绑定 OnHealthChanged 与 OnMaxHealthChanged。"],
        ["WBP_ManaGlobe", "Event WidgetControllerSet 中 Cast/Set 后绑定 OnManaChanged 与 OnMaxManaChanged。"],
    ],
    [1.8, 4.5],
)

doc.add_heading("10. 常见错误与定位", level=1)
add_table(
    doc,
    ["现象", "高概率原因", "检查位置"],
    [
        ["UI 完全不显示", "HUD 没被创建，或 OverlayWidgetClass 未设置。", "BP_AuraGameMode / BP_AuraHUD / AAuraHUD::InitOverlay"],
        ["OverlayWidget 空指针", "CreateWidget 后没有 Cast/赋值给 OverlayWidget。", "Source/Aura/Private/UI/HUD/AuraHUD.cpp"],
        ["蓝图报无访问 BPOverlayWidgetController", "Bind 事件时 Controller 变量还没 Set，或 Cast 失败。", "WBP_HealthGlobe / WBP_ManaGlobe Event WidgetControllerSet"],
        ["Mana 事件参数还叫 NewHealth", "OnManaChanged 使用了 Health 签名，或蓝图/UHT 缓存旧节点。", "OverlayWidgetController.h；删除旧蓝图节点并 Refresh All Nodes"],
        ["初始血量显示成最大血量", "MaxHealth 被广播到了 OnHealthChanged。", "OverlayWidgetController.cpp::BroadcastInitialValues"],
        ["属性变化但 UI 不动", "BindCallbacksToDependencies 没调用，或蓝图没有绑定事件。", "AAuraHUD::GetOverlayWidgetController / WBP 子控件"],
        ["Cast 到 BP_OverlayWidgetController 失败", "BP_AuraHUD 的 OverlayWidgetControllerClass 没设为 BP_OverlayWidgetController。", "BP_AuraHUD Details"],
    ],
    [1.8, 2.2, 2.3],
)

doc.add_heading("11. 推荐调试断点", level=1)
add_numbers(
    doc,
    [
        "Source/Aura/Private/Character/AuraCharacter.cpp：AAuraCharacter::InitAbilityActorInfo，确认是否进入 UI 初始化链。",
        "Source/Aura/Private/UI/HUD/AuraHUD.cpp：AAuraHUD::InitOverlay，确认 Widget 与 Controller 是否创建。",
        "Source/Aura/Private/UI/HUD/AuraHUD.cpp：AAuraHUD::GetOverlayWidgetController，确认 SetWidgetControllerParams 与 BindCallbacksToDependencies 是否执行。",
        "Source/Aura/Private/UI/Widget/AuraUserWidget.cpp：UAuraUserWidget::SetWidgetController，确认蓝图 Event WidgetControllerSet 的触发时机。",
        "Source/Aura/Private/UI/WidgetController/OverlayWidgetController.cpp：BroadcastInitialValues，确认初始值是否广播。",
        "Source/Aura/Private/UI/WidgetController/OverlayWidgetController.cpp：HealthChanged / ManaChanged，确认属性变化后是否进入回调。",
    ],
)

doc.add_heading("12. 心智模型", level=1)
add_code(
    doc,
    "AttributeSet: 真实属性账本，保存 Health / MaxHealth / Mana / MaxMana。\n"
    "AbilitySystemComponent: 事件源，属性变化时触发 Delegate。\n"
    "WidgetController: 数据中间层，监听 ASC，再广播给蓝图。\n"
    "HUD: 装配工，创建 Widget 和 WidgetController。\n"
    "Widget: 显示层，绑定 Controller 事件，更新进度条。\n"
    "WBP_Overlay: 主 UI，负责把 Controller 分发给子 Widget。\n"
    "WBP_HealthGlobe / WBP_ManaGlobe: 子 UI，负责具体显示 Health / Mana。",
)

doc.add_paragraph("最后记住这条最重要的时序：")
add_code(
    doc,
    "SetWidgetControllerParams\n"
    "    -> BindCallbacksToDependencies\n"
    "    -> OverlayWidget->SetWidgetController\n"
    "    -> 蓝图 WidgetControllerSet 绑定事件\n"
    "    -> BroadcastInitialValues\n"
    "    -> AddToViewport",
)

doc.core_properties.title = "Aura HUD WidgetController UI Guide"
doc.core_properties.subject = "Aura UE GAS UI architecture notes"
doc.core_properties.author = "Codex"

doc.save(OUT)
print(OUT)
