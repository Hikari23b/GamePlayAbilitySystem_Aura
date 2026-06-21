from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT = "Docs/Aura_GameplayEffect_Technical_Guide.docx"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Calibri"
    r.font.size = Pt(10)


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        cell_text(c, h, True)
        shade(c, "E8EEF5")
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if widths:
            c.width = Inches(widths[i])
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cell_text(cells[i], v)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths:
                cells[i].width = Inches(widths[i])
    doc.add_paragraph()


def code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(30, 30, 30)


def note(doc, text):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.rows[0].cells[0]
    shade(c, "F4F6F9")
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(10)
    doc.add_paragraph()


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(1)
sec.bottom_margin = Inches(1)
sec.left_margin = Inches(1)
sec.right_margin = Inches(1)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(11)
styles["Normal"].paragraph_format.space_after = Pt(6)
styles["Normal"].paragraph_format.line_spacing = 1.25

for name, size, color in [
    ("Heading 1", 16, "2E74B5"),
    ("Heading 2", 13, "2E74B5"),
    ("Heading 3", 12, "1F4D78"),
]:
    s = styles[name]
    s.font.name = "Calibri"
    s.font.size = Pt(size)
    s.font.color.rgb = RGBColor.from_string(color)
    s.paragraph_format.space_before = Pt(12)
    s.paragraph_format.space_after = Pt(6)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Aura GameplayEffect 技术文档")
r.bold = True
r.font.name = "Calibri"
r.font.size = Pt(20)
r.font.color.rgb = RGBColor.from_string("0B2545")

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run("范围：AuraEffectActor、GameplayEffect 应用/移除、AttributeSet 生命周期回调、EffectContext、Spec、Handle。")

note(doc, "核心记忆：GameplayEffect 是配置好的效果配方；ASC 负责应用；AttributeSet 保存属性；EffectActor 只是触发器。")

doc.add_heading("1. 总览", level=1)
doc.add_paragraph("这套 GameplayEffect 逻辑的目标是：让场景中的 EffectActor 在重叠开始或结束时，对目标 Actor 应用一个或多个 GameplayEffect，并能在离开范围时移除 Infinite 类型效果。")
code(
    doc,
    "Blueprint BeginOverlap\n"
    "    -> AAuraEffectActor::OnOverlap(TargetActor)\n"
    "        -> 遍历 Instant / Duration / Infinite 数组\n"
    "        -> ApplyEffectToTarget(TargetActor, EffectClass)\n"
    "            -> 获取 TargetASC\n"
    "            -> MakeEffectContext\n"
    "            -> MakeOutgoingSpec\n"
    "            -> ApplyGameplayEffectSpecToSelf\n"
    "            -> 若为 Infinite 且需要 EndOverlap 移除，保存 ActiveEffectHandle\n\n"
    "Blueprint EndOverlap\n"
    "    -> AAuraEffectActor::OnEndOverlap(TargetActor)\n"
    "        -> 按策略应用 EndOverlap Effects\n"
    "        -> 找到 TargetASC 对应的 Infinite Handles\n"
    "        -> RemoveActiveGameplayEffect\n"
    "        -> 从 ActiveEffectHandles 中清理记录",
)

doc.add_heading("2. 文件职责", level=1)
table(
    doc,
    ["文件", "核心内容", "职责"],
    [
        ["Source/Aura/Public/Actor/AuraEffectActor.h", "AAuraEffectActor、EEffectApplicationPolicy、EEffectRemovalPolicy", "声明 EffectActor 的应用策略、移除策略、Effect 数组、ActiveEffectHandles。"],
        ["Source/Aura/Private/Actor/AuraEffectActor.cpp", "ApplyEffectToTarget / OnOverlap / OnEndOverlap", "执行 GE 应用；保存 Infinite Handle；离开范围时移除对应 Infinite 效果。"],
        ["Source/Aura/Public/AbilitySystem/AuraAttributeSet.h", "UAuraAttributeSet、FEffectProperties", "声明 Health/Mana 属性，声明属性变化回调和效果属性上下文结构。"],
        ["Source/Aura/Private/AbilitySystem/AuraAttributeSet.cpp", "PreAttributeChange / PostGameplayEffectExecute / SetEffectProperties", "限制属性值范围；在 GE 执行后整理来源与目标信息。"],
        ["GameplayEffect 蓝图资产", "Modifiers / Duration Policy / Period / Stacking", "决定具体修改 Health、Mana 还是其他属性，以及效果持续/周期/堆叠行为。"],
        ["EffectActor 蓝图", "Mesh / Sphere / BeginOverlap / EndOverlap", "提供可视化和碰撞触发，并调用 C++ OnOverlap / OnEndOverlap。"],
    ],
    [2.1, 2.0, 2.2],
)

doc.add_heading("3. AuraEffectActor 结构", level=1)
doc.add_paragraph("文件：Source/Aura/Public/Actor/AuraEffectActor.h")
doc.add_heading("3.1 应用策略枚举", level=2)
code(
    doc,
    "UENUM(BlueprintType)\n"
    "enum class EEffectApplicationPolicy : uint8\n"
    "{\n"
    "    ApplyOnOverlap,\n"
    "    ApplyOnEndOverlap,\n"
    "    DoNotApply\n"
    "};",
)
bullets(
    doc,
    [
        "ApplyOnOverlap：进入碰撞范围时应用。",
        "ApplyOnEndOverlap：离开碰撞范围时应用。",
        "DoNotApply：不应用该类型效果。",
        "BlueprintType 的 enum class 必须指定 : uint8。",
    ],
)

doc.add_heading("3.2 移除策略枚举", level=2)
code(
    doc,
    "UENUM(BlueprintType)\n"
    "enum class EEffectRemovalPolicy : uint8\n"
    "{\n"
    "    RemoveOnEndOverlap,\n"
    "    DoNotRemove\n"
    "};",
)
bullets(
    doc,
    [
        "RemoveOnEndOverlap：离开范围时移除之前保存的 Infinite Effects。",
        "DoNotRemove：离开范围后不移除，Infinite 效果会继续留在目标 ASC 上。",
    ],
)

doc.add_heading("3.3 Effect 数组", level=2)
code(
    doc,
    "UPROPERTY(EditAnywhere, BlueprintReadOnly, Category = \"Applied Effects\")\n"
    "TArray<TSubclassOf<UGameplayEffect>> InstantGameplayEffectClasses;\n\n"
    "UPROPERTY(EditAnywhere, BlueprintReadOnly, Category = \"Applied Effects\")\n"
    "TArray<TSubclassOf<UGameplayEffect>> DurationGameplayEffectClasses;\n\n"
    "UPROPERTY(EditAnywhere, BlueprintReadOnly, Category = \"Applied Effects\")\n"
    "TArray<TSubclassOf<UGameplayEffect>> InfiniteGameplayEffectClasses;",
)
bullets(
    doc,
    [
        "数组版本允许一个 EffectActor 同时应用多个 GE。",
        "蓝图 Event Graph 不需要拖数组出来；只需要 BeginOverlap 调 OnOverlap，EndOverlap 调 OnEndOverlap。",
        "具体 GE 在 EffectActor 蓝图 Details 面板里配置。",
    ],
)
note(doc, "当前代码里仍留有单个 InstantGameplayEffectClass / DurationGameplayEffectClass / InfiniteGameplayEffectClass 字段，但已不参与数组版逻辑。后续可以清理，避免蓝图配置混淆。")

doc.add_heading("4. ApplyEffectToTarget", level=1)
doc.add_paragraph("文件：Source/Aura/Private/Actor/AuraEffectActor.cpp")
doc.add_paragraph("函数：AAuraEffectActor::ApplyEffectToTarget(AActor* TargetActor, TSubclassOf<UGameplayEffect> GameplayEffectClass)")
code(
    doc,
    "UAbilitySystemComponent* TargetASC = UAbilitySystemBlueprintLibrary::GetAbilitySystemComponent(TargetActor);\n"
    "if (TargetASC == nullptr) return;\n"
    "if (!GameplayEffectClass) return;\n\n"
    "FGameplayEffectContextHandle EffectContextHandle = TargetASC->MakeEffectContext();\n"
    "EffectContextHandle.AddSourceObject(this);\n\n"
    "FGameplayEffectSpecHandle EffectSpecHandle = TargetASC->MakeOutgoingSpec(GameplayEffectClass, ActorLevel, EffectContextHandle);\n"
    "FActiveGameplayEffectHandle ActiveEffectHandle = TargetASC->ApplyGameplayEffectSpecToSelf(*EffectSpecHandle.Data.Get());",
)
table(
    doc,
    ["步骤", "含义"],
    [
        ["GetAbilitySystemComponent", "从目标 Actor 上获取 ASC。没有 ASC 就无法应用 GE。"],
        ["MakeEffectContext", "创建这次效果的上下文，记录来源、Instigator、SourceObject 等背景信息。"],
        ["AddSourceObject(this)", "把当前 AuraEffectActor 记录为本次效果来源对象。"],
        ["MakeOutgoingSpec", "根据 GE 类、等级 ActorLevel、上下文创建这一次具体效果实例。"],
        ["ApplyGameplayEffectSpecToSelf", "把 Spec 应用到目标自己的 ASC 上，返回 ActiveEffectHandle。"],
    ],
    [2.0, 4.3],
)

doc.add_heading("5. Context / Spec / Handle", level=1)
table(
    doc,
    ["概念", "可以这样理解", "作用"],
    [
        ["GameplayEffectClass", "配方", "决定修改哪些属性、持续多久、是否周期、是否堆叠。"],
        ["EffectContextHandle", "现场记录", "记录来源对象、Instigator、命中信息、上下文来源等。"],
        ["EffectSpecHandle", "本次执行单据", "由 GE 配方 + 等级 + Context 生成的具体效果实例。"],
        ["ActiveGameplayEffectHandle", "已激活效果的编号", "应用成功后用于以后查找或移除这个 Active GE。"],
        ["AbilitySystemComponent", "账本管理器", "保存 ActiveGameplayEffects，并负责应用、复制、移除 GE。"],
    ],
    [1.5, 1.6, 3.2],
)
note(doc, "FActiveGameplayEffectHandle 不是效果本体。真正的 Active GameplayEffect 数据在 ASC 内部；Handle 只是钥匙。")

doc.add_heading("6. Infinite 效果保存与移除", level=1)
doc.add_heading("6.1 保存 Handle", level=2)
code(
    doc,
    "bool bIsInfinite = EffectSpecHandle.Data.Get()->Def.Get()->DurationPolicy == EGameplayEffectDurationType::Infinite;\n\n"
    "if (bIsInfinite && InfiniteEffectRemovalPolicy == EEffectRemovalPolicy::RemoveOnEndOverlap)\n"
    "{\n"
    "    ActiveEffectHandles.Add(ActiveEffectHandle, TargetASC);\n"
    "}",
)
bullets(
    doc,
    [
        "只有 Infinite GE 需要手动移除，所以只有 Infinite 才保存 Handle。",
        "ActiveEffectHandles 的 Key 是 Handle，Value 是拥有该效果的 ASC。",
        "保存 ASC 是因为 Handle 必须交给对应 ASC 才能移除。",
    ],
)

doc.add_heading("6.2 移除 Handle", level=2)
code(
    doc,
    "if (InfiniteEffectRemovalPolicy == EEffectRemovalPolicy::RemoveOnEndOverlap)\n"
    "{\n"
    "    UAbilitySystemComponent* TargetASC = UAbilitySystemBlueprintLibrary::GetAbilitySystemComponent(TargetActor);\n"
    "    if (!IsValid(TargetASC)) return;\n\n"
    "    TArray<FActiveGameplayEffectHandle> HandlesToRemove;\n"
    "    for (const auto& HandlePair : ActiveEffectHandles)\n"
    "    {\n"
    "        if (TargetASC == HandlePair.Value)\n"
    "        {\n"
    "            TargetASC->RemoveActiveGameplayEffect(HandlePair.Key, 1);\n"
    "            HandlesToRemove.Add(HandlePair.Key);\n"
    "        }\n"
    "    }\n\n"
    "    for (FActiveGameplayEffectHandle& Handle : HandlesToRemove)\n"
    "    {\n"
    "        ActiveEffectHandles.FindAndRemoveChecked(Handle);\n"
    "    }\n"
    "}",
)
bullets(
    doc,
    [
        "先遍历 Map，找到属于当前 TargetASC 的所有 Infinite Handles。",
        "调用 RemoveActiveGameplayEffect 移除目标身上的 Active GE。",
        "不要遍历 Map 时直接删除 Map 元素；先记录到 HandlesToRemove，循环结束后再删。",
        "从 TMap 移除 Key 会同时移除对应 Value 记录，但不会销毁 ASC 对象。",
    ],
)
note(doc, "如果一个 Infinite GE 可能叠多层，RemoveActiveGameplayEffect 的 StacksToRemove 参数会影响移除几层。传 1 通常只移除一层；若想清掉整个 Active GE，需要按当前教程/引擎版本确认是否使用默认值或 -1。")

doc.add_heading("7. OnOverlap 与 OnEndOverlap", level=1)
doc.add_paragraph("蓝图侧 BeginOverlap 不再直接调用 ApplyEffectToTarget，而是调用 OnOverlap。C++ 内部根据策略遍历数组。")
code(
    doc,
    "void AAuraEffectActor::OnOverlap(AActor* TargetActor)\n"
    "{\n"
    "    if (InstantEffectApplicationPolicy == EEffectApplicationPolicy::ApplyOnOverlap)\n"
    "    {\n"
    "        for (TSubclassOf<UGameplayEffect> EffectClass : InstantGameplayEffectClasses)\n"
    "        {\n"
    "            ApplyEffectToTarget(TargetActor, EffectClass);\n"
    "        }\n"
    "    }\n"
    "    // Duration / Infinite 同理\n"
    "}",
)
bullets(
    doc,
    [
        "蓝图只传 OtherActor，不需要传具体 GE Class。",
        "GE Class 在 Actor 蓝图 Details 面板里的数组里配置。",
        "OnEndOverlap 既可以按策略应用 EndOverlap Effects，也负责移除保存过的 Infinite Effects。",
    ],
)

doc.add_heading("8. GameplayEffect 蓝图配置", level=1)
table(
    doc,
    ["目标", "Duration Policy", "常见配置", "说明"],
    [
        ["瞬间回血/回蓝", "Instant", "Modifier: Health/Mana Add 25", "立即执行，执行后结束，不需要 Handle。"],
        ["持续 5 秒 Buff", "Has Duration", "Duration=5，Modifier 或周期执行", "到时间自动结束，通常不需要手动移除。"],
        ["区域内持续扣血", "Infinite", "Period=1，Modifier 或 Execution", "不会自动结束，离开区域时需要 RemoveActiveGameplayEffect。"],
        ["站在光环内加速", "Infinite", "MoveSpeed Add/Multiply", "进入应用，离开移除。"],
    ],
    [1.4, 1.2, 2.3, 1.7],
)

doc.add_heading("9. AttributeSet 回调", level=1)
doc.add_paragraph("文件：Source/Aura/Private/AbilitySystem/AuraAttributeSet.cpp")

doc.add_heading("9.1 PreAttributeChange", level=2)
code(
    doc,
    "void UAuraAttributeSet::PreAttributeChange(const FGameplayAttribute& Attribute, float& NewValue)\n"
    "{\n"
    "    Super::PreAttributeChange(Attribute, NewValue);\n"
    "    if (Attribute == GetHealthAttribute())\n"
    "    {\n"
    "        NewValue = FMath::Clamp(NewValue, 0.f, GetMaxHealth());\n"
    "    }\n"
    "    if (Attribute == GetManaAttribute())\n"
    "    {\n"
    "        NewValue = FMath::Clamp(NewValue, 0.f, GetMaxMana());\n"
    "    }\n"
    "}",
)
bullets(
    doc,
    [
        "这是 GAS 生命周期回调，不需要手动调用。",
        "Attribute 表示即将改变哪个属性。",
        "NewValue 是引用，修改它会影响最终写入的值。",
        "当前逻辑确保 Health/Mana 不低于 0，也不超过最大值。",
    ],
)

doc.add_heading("9.2 PostGameplayEffectExecute", level=2)
doc.add_paragraph("GE 修改属性后，GAS 会调用 PostGameplayEffectExecute。它适合做伤害结算、死亡判断、飘字、经验归属等后处理。")
code(
    doc,
    "void UAuraAttributeSet::PostGameplayEffectExecute(const FGameplayEffectModCallbackData& Data)\n"
    "{\n"
    "    Super::PostGameplayEffectExecute(Data);\n"
    "    FEffectProperties Props;\n"
    "    SetEffectProperties(Data, Props);\n"
    "}",
)

doc.add_heading("9.3 FEffectProperties 与 SetEffectProperties", level=2)
doc.add_paragraph("FEffectProperties 用来把 Data 里的来源/目标信息整理成一个容易使用的结构。")
table(
    doc,
    ["字段", "含义"],
    [
        ["EffectContextHandle", "本次 GE 的上下文。"],
        ["SourceASC", "效果来源方 ASC。"],
        ["SourceAvatarActor", "来源 ASC 的 AvatarActor。"],
        ["SourceController", "来源控制器，玩家可能是 PlayerController，AI 可能是 AIController。"],
        ["SourceCharacter", "来源角色。"],
        ["TargetASC", "被影响目标的 ASC。"],
        ["TargetAvatarActor", "被影响目标的 AvatarActor。"],
        ["TargetController", "目标控制器。"],
        ["TargetCharacter", "目标角色。"],
    ],
    [2.0, 4.3],
)
code(
    doc,
    "Props.EffectContextHandle = Data.EffectSpec.GetContext();\n"
    "Props.SourceASC = Props.EffectContextHandle.GetOriginalInstigatorAbilitySystemComponent();\n"
    "Props.SourceAvatarActor = Props.SourceASC->AbilityActorInfo->AvatarActor.Get();\n"
    "Props.TargetAvatarActor = Data.Target.AbilityActorInfo->AvatarActor.Get();\n"
    "Props.TargetASC = UAbilitySystemBlueprintLibrary::GetAbilitySystemComponent(Props.TargetAvatarActor);",
)
note(doc, "SetEffectProperties 的目的不是改变属性，而是整理信息。后面处理伤害、死亡、击杀者、飘字时会大量用 Props.Source... 和 Props.Target...。")

doc.add_heading("10. 蓝图侧推荐接法", level=1)
numbers(
    doc,
    [
        "EffectActor 蓝图继承 AAuraEffectActor。",
        "添加 StaticMesh / Sphere 等组件，Sphere 负责重叠事件。",
        "BeginOverlap(Sphere) 的 OtherActor 连接到 OnOverlap(TargetActor)。",
        "EndOverlap(Sphere) 的 OtherActor 连接到 OnEndOverlap(TargetActor)。",
        "Details 面板中配置 Instant / Duration / Infinite GameplayEffectClasses 数组。",
        "区域类 Infinite 效果：InfiniteEffectApplicationPolicy = ApplyOnOverlap，InfiniteEffectRemovalPolicy = RemoveOnEndOverlap。",
    ],
)

doc.add_heading("11. 常见问题排查", level=1)
table(
    doc,
    ["现象", "高概率原因", "检查位置"],
    [
        ["进入区域有效，离开不移除", "EndOverlap 没调用 OnEndOverlap；或 ActiveEffectHandles 没保存；或 TargetASC 不匹配。", "BP EventGraph / AuraEffectActor.cpp"],
        ["ActiveEffectHandles.Add 没执行", "GE 不是 Infinite，或 RemovalPolicy 不是 RemoveOnEndOverlap。", "GE Duration Policy / EffectActor Details"],
        ["MapBefore/MapAfter 都是 1", "TMap 可能覆盖同 Key，或已有一条记录；不代表 Add 没执行。", "UE_LOG 输出"],
        ["ApplyEffectToTarget 没效果", "TargetActor 没 ASC，或 GameplayEffectClass 为空。", "GetAbilitySystemComponent / Details 数组"],
        ["血量超过最大值", "PreAttributeChange 没重写或没有 Clamp Health。", "AuraAttributeSet.cpp"],
        ["蓝图下拉能看到 AuraAttributeSet.Health", "这是 UPROPERTY + FGameplayAttributeData + UAttributeSet 反射结果。", "AuraAttributeSet.h"],
        ["UENUM BlueprintType 报错", "enum class 没指定 : uint8。", "AuraEffectActor.h"],
        ["TMap Handle 编译报 GetTypeHash", "头文件没有 include GameplayEffectTypes.h。", "AuraEffectActor.h"],
    ],
    [1.8, 2.5, 2.0],
)

doc.add_heading("12. 推荐调试日志", level=1)
code(
    doc,
    "UE_LOG(LogTemp, Warning, TEXT(\"bIsInfinite: %d, RemovalPolicy: %d, HandleValid: %d, MapBefore: %d\"),\n"
    "    bIsInfinite,\n"
    "    static_cast<int32>(InfiniteEffectRemovalPolicy),\n"
    "    ActiveEffectHandle.IsValid(),\n"
    "    ActiveEffectHandles.Num());\n\n"
    "ActiveEffectHandles.Add(ActiveEffectHandle, TargetASC);\n"
    "UE_LOG(LogTemp, Warning, TEXT(\"Added Handle. MapAfter: %d\"), ActiveEffectHandles.Num());",
)
code(
    doc,
    "UE_LOG(LogTemp, Warning, TEXT(\"OnEndOverlap called. MapNum: %d\"), ActiveEffectHandles.Num());\n"
    "UE_LOG(LogTemp, Warning, TEXT(\"SameASC: %d, HandleValid: %d\"),\n"
    "    TargetASC == HandlePair.Value,\n"
    "    HandlePair.Key.IsValid());",
)

doc.add_heading("13. 心智模型", level=1)
code(
    doc,
    "GameplayEffect = 配方：定义改什么属性、怎么改、持续多久、是否周期/堆叠。\n"
    "AuraEffectActor = 触发器：进入/离开范围时应用或移除 GE。\n"
    "AbilitySystemComponent = 执行者和管理者：应用 GE，保存 Active GE，负责移除。\n"
    "GameplayEffectSpec = 本次效果实例：GE 类 + 等级 + Context。\n"
    "GameplayEffectContext = 现场信息：来源对象、Instigator、HitResult 等。\n"
    "ActiveGameplayEffectHandle = 钥匙：用于之后找到并移除已应用效果。\n"
    "AttributeSet = 属性账本：保存 Health/Mana，并在回调里限制或结算属性。",
)

doc.core_properties.title = "Aura GameplayEffect Technical Guide"
doc.core_properties.subject = "Aura UE GAS GameplayEffect notes"
doc.core_properties.author = "Codex"
doc.save(OUT)
print(OUT)
